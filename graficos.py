"""
Analizador Dinámico de Plataformas de Soporte para Equipos Rotativos
======================================================================
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io, math, warnings, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as q
warnings.filterwarnings("ignore")

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Analizador Dinámico — Plataformas de Soporte",
    page_icon="📊", layout="wide", initial_sidebar_state="expanded"
)

# Se eliminaron los colores forzados para garantizar compatibilidad 100% con Dark/Light Mode nativo.
st.markdown("""
<style>
.stDataFrame { font-size: 12px; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTES Y ZONAS DE CLASIFICACIÓN
# ══════════════════════════════════════════════════════════════════════════════
F_REF_N = 9810.0  # N/Ton (fuerza de referencia SAP2000)

RICHART_ZONES = [
    (0.787,  "No perceptible",          "#1a9850"),
    (1.854,  "Apenas perceptible",      "#66bd63"),
    (5.309,  "Fácilmente perceptible",  "#fee08b"),
    (10.640, "Molesta",                 "#fdae61"),
    (21.290, "Severa (personas)",       "#f46d43"),
    (53.090, "Límite máquinas",         "#d73027"),
    (127.00, "Precaución estructuras",  "#9e0142"),
    (1e9,    "Peligro estructuras",     "#67001f"),
]
BLAKE_ZONES = [
    (0.360,  "A — Sin fallas",         "#1a9850"),
    (3.590,  "B — Fallas menores",     "#66bd63"),
    (35.90,  "C — Defectuoso",         "#fee08b"),
    (107.7,  "D — Falla inminente",    "#f46d43"),
    (1e9,    "E — Peligroso",          "#d73027"),
]
ISO_ZONES = [
    (1.12, "Zona A — Aceptable",        "#1a9850"),
    (2.80, "Zona B — Normal",           "#66bd63"),
    (7.10, "Zona C — Alarma",           "#fee08b"),
    (1e9,  "Zona D — Detener equipo",   "#f46d43"),
]

# Curvas de los diagramas de clasificación (x en cpm, y en mm). Se usan tanto en
# las gráficas interactivas (Plotly) como en las figuras del reporte (matplotlib).
RICHART_LINES = {
    "Peligro estructuras":    {"x": [100, 580, 4100, 10000], "y": [10.08, 1.8796, 0.2921, 0.1250], "col": "#67001f", "dash": "solid"},
    "Precaución estructuras": {"x": [100, 640, 3500, 10000], "y": [9.113, 1.3462, 0.2311, 0.0785], "col": "#9e0142", "dash": "solid"},
    "Límite máquinas":        {"x": [100, 4600, 10000], "y": [2.54, 0.05461, 0.02506], "col": "#d73027", "dash": "solid"},
    "Severa (personas)":      {"x": [100, 540, 2600, 4200, 10000], "y": [14.46, 0.4572, 0.01829, 0.01067, 0.00400], "col": "#f46d43", "dash": "dash"},
    "Molesta (personas)":     {"x": [100, 270, 2250, 3800, 10000], "y": [2.846, 0.5588, 0.01778, 0.00965, 0.00312], "col": "#fdae61", "dash": "dash"},
    "Fácilmente perceptible": {"x": [100, 260, 5700, 10000], "y": [0.268, 0.10287, 0.00465, 0.00264], "col": "#fee08b", "dash": "dash"},
    "Apenas perceptible":     {"x": [100, 185, 3900, 10000], "y": [0.0846, 0.04699, 0.00254, 0.00103], "col": "#66bd63", "dash": "dash"},
    "No perceptible":         {"x": [100, 145, 1200, 10000], "y": [0.0336, 0.02286, 0.00254, 0.00028], "col": "#1a9850", "dash": "dash"},
}
BLAKE_LINES = {
    "E — Peligroso":       {"x": [100, 310, 1850, 3700, 10000], "y": [0.5842, 1.016, 0.2286, 0.12192, 0.01524], "col": "#d73027", "dash": "solid"},
    "D — Falla inminente": {"x": [100, 310, 1850, 3700, 10000], "y": [0.2286, 0.18288, 0.08636, 0.04191, 0.004826], "col": "#f46d43", "dash": "solid"},
    "C — Defectuoso":      {"x": [100, 310, 1850, 3700, 10000], "y": [0.09398, 0.0635, 0.02921, 0.013208, 0.00381], "col": "#fee08b", "dash": "solid"},
    "B — Fallas menores":  {"x": [100, 1850, 5900, 10000], "y": [0.0381, 0.009017, 0.00254, 0.001417], "col": "#66bd63", "dash": "solid"},
}

def classify(v, zones):
    for v_max, label, color in zones:
        if v <= v_max:
            return label, color
    return zones[-1][1], zones[-1][2]

def v_peak(amp_mm, f_hz):
    return amp_mm * 2 * math.pi * f_hz

def v_rms(amp_mm, f_hz):
    return v_peak(amp_mm, f_hz) / math.sqrt(2)

def auto_dir(caso):
    """Detección automática de columna de respuesta y dirección desde el nombre del caso."""
    caso_u = str(caso).upper()
    if '_X' in caso_u or caso_u.endswith('X'): return ('U1', 'X')
    if '_Y' in caso_u or caso_u.endswith('Y'): return ('U2', 'Y')
    if '_Z' in caso_u or caso_u.endswith('Z'): return ('U3', 'Z')
    return ('U1', str(caso))

def _fuerza_tr(modo_fuerza, f_arr, F_rd_cal, U_gmm, n_apoyos):
    """Fuerza dinámica del transitorio en cada frecuencia [N].
    - Valor fijo:      F_rd del fabricante, constante (ya viene por apoyo).
    - Curva desbalance: F=m·e·ω² (TOTAL del rotor) repartido entre n_apoyos."""
    if modo_fuerza.startswith("Curva"):
        return (U_gmm / 1e6) * (2 * np.pi * f_arr) ** 2 / max(int(n_apoyos), 1)
    return np.full(np.asarray(f_arr, float).shape, float(F_rd_cal))


def transitorio_uf(r, F_rd_cal, modo_fuerza, U_gmm, F_REF, f_tr_lo=None, f_tr_hi=None, n_apoyos=1):
    """Respuesta del transitorio (partida/parada) — PEOR caso del tramo.

    Recorre la ventana de barrido [f_tr_lo, f_tr_hi] y devuelve el PEOR caso de
    amplitud u(f) = |H(f)|·F(f)/F_REF — no solo el punto f_rd.
    - Valor fijo:      F(f) = F_rd constante en toda la ventana (conservador). El
      máximo de u(f) coincide entonces con el máximo de |H(f)|.
    - Curva desbalance: F(f) = m·e·ω²/n_apoyos. El máximo de u(f) pondera FRF y
      fuerza (la fuerza crece con f²).
    Si la ventana no solapa con los datos, cae al punto más cercano a f_rd.

    Devuelve (amplitud_mm, f_peor[Hz], FRF_usada[mm/u.carga], F_usada[N]),
    todo evaluado en la frecuencia del peor caso DENTRO de este nodo/caso.
    """
    freqs = np.array(r['freqs']); frf = np.array(r['frf_mm'])
    if freqs.size == 0:
        return 0.0, r['f_rd'], 0.0, float(F_rd_cal)
    lo = freqs.min() if f_tr_lo is None else min(f_tr_lo, f_tr_hi)
    hi = freqs.max() if f_tr_hi is None else max(f_tr_lo, f_tr_hi)
    win = (freqs >= lo) & (freqs <= hi)
    if not win.any():
        # Ventana sin datos: cae al punto más cercano a f_rd
        idx0 = int(np.argmin(np.abs(freqs - r['f_rd'])))
        win = np.zeros(len(freqs), bool); win[idx0] = True
    fw = freqs[win]; hw = frf[win]
    F_arr = _fuerza_tr(modo_fuerza, fw, F_rd_cal, U_gmm, n_apoyos)
    u_arr = hw * F_arr / F_REF
    idx = int(np.argmax(u_arr))
    return float(u_arr[idx]), float(fw[idx]), float(hw[idx]), float(F_arr[idx])


def transitorio_en_frd(r, F_rd_cal, modo_fuerza, U_gmm, F_REF, n_apoyos=1):
    """Respuesta evaluada EN EL PUNTO de cruce del aislador f_rd (no el peak del
    tramo). Devuelve (amplitud_mm, f_usada[Hz], F_usada[N], en_rango[bool]). Si f_rd
    cae fuera de los datos, usa el punto más cercano y marca en_rango=False.
    """
    freqs = np.array(r['freqs']); frf = np.array(r['frf_mm'])
    if freqs.size == 0:
        return float('nan'), r['f_rd'], float(F_rd_cal), False
    f_rd = r['f_rd']
    en_rango = bool(freqs.min() <= f_rd <= freqs.max())
    idx = int(np.argmin(np.abs(freqs - f_rd)))
    f_used = float(freqs[idx]); h = float(frf[idx])
    F = float(_fuerza_tr(modo_fuerza, np.array([f_used]), F_rd_cal, U_gmm, n_apoyos)[0])
    return h * F / F_REF, f_used, F, en_rango


def fig_clasif_loglog_mpl(lines, puntos, f_excl_lo, f_excl_hi, f_op, titulo, ylim):
    """Diagrama log-log de clasificación (Richart/Blake) en matplotlib, para el Word.
    - lines: {label: {'x':[cpm], 'y':[mm], 'col':hex, 'dash':'solid'|'dash'}}
    - puntos: lista de (f_cpm, A_mm, etiqueta, color_hex, marker)  ['o'=Op, '^'=Tr]
    """
    fig, ax = plt.subplots(figsize=(9, 6.2))
    for lbl, d in lines.items():
        ax.plot(d['x'], d['y'], color=d['col'], lw=1.6,
                ls='--' if d.get('dash') == 'dash' else '-', label=lbl)
    ax.axvspan(f_excl_lo*60, f_excl_hi*60, color='orange', alpha=0.15)
    ax.axvline(f_op*60, color='red', ls='--', lw=1.3, alpha=0.8)
    for f_cpm, A, lbl, col_, mk in puntos:
        if A is None or not np.isfinite(A) or A <= 0:
            continue
        ax.plot(f_cpm, A, marker=mk, ms=9, color=col_,
                markeredgecolor='black', markeredgewidth=1.0, ls='none', zorder=5)
        ax.annotate(lbl, (f_cpm, A), fontsize=6, xytext=(4, 3), textcoords='offset points')
    ax.set_xscale('log'); ax.set_yscale('log')
    ax.set_xlim(100, 10000); ax.set_ylim(*ylim)
    ax.set_xlabel('Frecuencia (cpm)', fontsize=9)
    ax.set_ylabel('Amplitud peak (mm)', fontsize=9)
    ax.set_title(titulo, fontsize=11, fontweight='bold')
    ax.grid(True, which='both', alpha=0.25, ls='--')
    ax.legend(fontsize=6.5, ncol=2, loc='lower left')
    return fig


def fig_iso_barras_mpl(df, iso_a, iso_b, iso_c, titulo):
    """Barras de v_RMS (operación) por nodo con bandas ISO 20816-3, para el Word."""
    fig, ax = plt.subplots(figsize=(10, 5))
    labels = (df['Joint'].astype(str) + "\n(" + df['Caso'].astype(str) + ")").tolist()
    vals = df['vRMS_op (mm/s)'].tolist()
    top = max((max(vals) * 1.2 if vals else 0), iso_c * 1.3, 10.0)
    for y0, y1, c in [(0, iso_a, '#1a9850'), (iso_a, iso_b, '#66bd63'),
                      (iso_b, iso_c, '#fee08b'), (iso_c, top, '#f46d43')]:
        ax.axhspan(y0, y1, color=c, alpha=0.12)
    for y, c, n in [(iso_a, '#1a9850', f'A/B = {iso_a:.2f}'),
                    (iso_b, '#b8860b', f'B/C = {iso_b:.2f}'),
                    (iso_c, '#f46d43', f'C/D = {iso_c:.2f}')]:
        ax.axhline(y, color=c, ls='--', lw=1.3)
        ax.text(len(labels)-0.4, y, n, fontsize=7, color=c, va='bottom', ha='right')
    ax.bar(range(len(labels)), vals, color='#1f77b4', alpha=0.85)
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels, fontsize=7)
    ax.set_ylim(0, top); ax.set_ylabel('v_RMS (mm/s)', fontsize=9)
    ax.set_title(titulo, fontsize=11, fontweight='bold')
    ax.grid(True, axis='y', alpha=0.25, ls='--')
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR — PARÁMETROS
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Parámetros del Proyecto")
    st.divider()

    st.markdown("**Equipo**")
    f_op_rpm = st.number_input("Frecuencia de operación (RPM)", value=2900.0, step=10.0, format="%.0f")
    f_op = f_op_rpm / 60.0
    st.caption(f"= {f_op:.2f} Hz (máxima del equipo)")
    W_pata = st.number_input("Peso del equipo por pata (kgf)", value=0.0, step=100.0, format="%.0f",
                             help="Peso estático por apoyo, en kgf (= masa en kg). Necesario para "
                                  "calcular f_rd del aislador. Si lo dejas en 0, ingresa f_rd manual.")

    st.divider()
    st.markdown("**Cargas del fabricante (N/pata = total ÷ 4)**")
    col1, col2 = st.columns(2)
    with col1:
        F_op_V  = st.number_input("F_op vertical",     value=425,  step=50)
        F_op_H  = st.number_input("F_op horizontal",   value=425,  step=50)
    with col2:
        F_rd_V  = st.number_input("F_rd vertical",     value=5000, step=100)
        F_rd_H  = st.number_input("F_rd horizontal",   value=2500, step=100)

    st.divider()
    st.markdown("**Aislador**")
    K_din = st.number_input("K_din aislador (N/mm)", value=5600.0, step=100.0, format="%.0f")
    st.caption("Novibra RAEM2500: 5600 N/mm")

    st.divider()
    st.markdown("**Transitorio (partida / parada)**")
    # Frecuencia del modo de cuerpo rígido del equipo sobre el aislador:
    #   f_rd = (1/2π)·√(k/m)  con  k = K_din[N/m]  y  m = W_pata[kgf] = masa[kg]
    if W_pata > 0:
        f_rd_calc = (1.0/(2*math.pi)) * math.sqrt(K_din*1000.0 / W_pata)
        st.caption(f"f_rd calculada (aislador): **{f_rd_calc:.2f} Hz**")
    else:
        f_rd_calc = None
        st.caption("Ingresa el peso por pata para calcular f_rd automáticamente.")
    _help_frd = ("f_rd es la frecuencia del **modo de cuerpo rígido** del equipo sobre el "
                 "aislador:\n\n**f_rd = (1/2π)·√(K_din / m)**\n\n"
                 "con K_din = rigidez dinámica del aislador y m = masa por apoyo. "
                 "Corresponde a la frecuencia que el equipo **atraviesa al partir o detenerse** "
                 "(condición transitoria de partida/parada): al subir o bajar de RPM, la "
                 "máquina cruza esta resonancia.\n\n"
                 "• **Calculada**: se obtiene de la expresión con el peso por pata.\n"
                 "• **Manual**: debe ser **otorgada por el fabricante** del aislador.")
    origen_frd = st.radio("Origen de f_rd", ["Calculada", "Manual"],
                          index=1 if f_rd_calc is None else 0, horizontal=True,
                          help=_help_frd)
    if origen_frd == "Calculada" and f_rd_calc is not None:
        f_rd = f_rd_calc
    else:
        f_rd = st.number_input("f_rd manual (Hz)", value=3.5, step=0.1, format="%.2f",
                               help="Valor de f_rd (frecuencia de partida/parada) **otorgado "
                                    "por el fabricante** del aislador, en Hz.")
    st.caption(f"f_rd en uso: **{f_rd:.2f} Hz**")

    modo_fuerza_tr = st.selectbox("Modelo de fuerza transitoria",
                                  ["Valor fijo (F_rd)", "Curva desbalance m·e·ω²"],
                                  help="**Valor fijo (F_rd)**: fuerza dinámica **constante** que "
                                       "se considera para efectos prácticos en la condición "
                                       "transitoria (partida/parada), como en los casos revisados "
                                       "en este proyecto. Se aplica la misma F_rd en toda la "
                                       "ventana del barrido.\n\n"
                                       "**Curva desbalance (m·e·ω²)**: la fuerza varía con f² "
                                       "(F = m·e·ω²); mínima en f_rd y máxima cerca de operación. "
                                       "Sólo aplica a equipos centrífugos / rotativos.")
    if modo_fuerza_tr.startswith("Curva"):
        st.info("Modelo de **desbalance rotativo** (F = m·e·ω²): aplica a equipos "
                "**centrífugos / rotativos** (bombas, ventiladores, sopladores, compresores, "
                "motores). No representa máquinas recíprocas ni de impacto.", icon="ℹ️")
        metodo_U = st.radio("Definir desbalance por", ["Valor U directo", "Grado ISO 1940 (G)"],
                            horizontal=True)
        if metodo_U.startswith("Grado"):
            m_rotor = st.number_input("Masa del rotor (kg)", value=0.0, step=10.0, format="%.1f",
                                      help="Masa de la parte ROTANTE (no el peso total del equipo).")
            G_grade = st.selectbox("Grado de balanceo G — ISO 1940-1 (mm/s)",
                                   [0.4, 1.0, 2.5, 6.3, 16.0, 40.0], index=3,
                                   help="G2.5: turbinas/compresores. G6.3: bombas y ventiladores "
                                        "(típico). G16: motores diésel/ejes de transmisión.")
            omega = 2*math.pi*f_op
            U_gmm = (m_rotor * 1000.0 * G_grade / omega) if omega > 0 else 0.0
            st.caption(f"U = m·1000·G/ω = **{U_gmm:,.1f} g·mm** (G{G_grade}, f_op = {f_op:.1f} Hz)")
        else:
            U_gmm = st.number_input("Desbalance U = m·e (g·mm)", value=0.0, step=10.0, format="%.1f",
                                    help="U = m·e del ROTOR: m = masa de la parte ROTANTE (NO el peso "
                                         "total del equipo), e = excentricidad. F(f)=m·e·ω².")
        n_apoyos = st.number_input("N° de apoyos (reparto por pata)", value=4, min_value=1, step=1,
                                   help="La FRF de SAP está normalizada por apoyo (1 ton/pata), así que "
                                        "la fuerza de desbalance total se divide entre los apoyos.")
        if U_gmm > 0:
            F_op_total = (U_gmm/1e6)*(2*math.pi*f_op)**2
            st.caption(f"A f_op: F_total = {F_op_total:.0f} N → {F_op_total/n_apoyos:.0f} N por apoyo")
        st.caption("Ref.: ISO 1940-1; Arya, O'Neill & Pincus (1979); Den Hartog, *Mechanical "
                   "Vibrations*. **Definición referencial** — reemplazar por la curva "
                   "fuerza–frecuencia del fabricante cuando esté disponible.")
    else:
        U_gmm = 0.0
        n_apoyos = 1

    st.markdown("**Ventana del transitorio (barrido)**")
    auto_win = st.checkbox("Auto: f_rd → f_op", value=True,
                           help="Tramo que el equipo recorre en partida/parada. "
                                "El peak transitorio se busca aquí, no solo en f_rd.")
    if auto_win:
        f_tr_lo, f_tr_hi = min(f_rd, f_op), max(f_rd, f_op)
        st.caption(f"Ventana = {f_tr_lo:.2f} – {f_tr_hi:.2f} Hz")
    else:
        col_pb1, col_pb2 = st.columns(2)
        with col_pb1:
            f_tr_lo = st.number_input("f inf (Hz)", value=float(round(min(f_rd, f_op), 2)),
                                      step=0.5, format="%.2f")
        with col_pb2:
            f_tr_hi = st.number_input("f sup (Hz)", value=float(round(max(f_rd, f_op), 2)),
                                      step=0.5, format="%.2f")
    st.caption("Se busca el peor |H(f)|·F(f) en este tramo. Puedes abarcar también la "
               "zona de exclusión si es relevante. El peak de **operación** se evalúa "
               "aparte, dentro de la zona de exclusión.")

    st.divider()
    st.markdown("**Zona de exclusión**")
    z_lo = st.number_input("Factor inferior", value=0.8, step=0.05, format="%.2f")
    z_hi = st.number_input("Factor superior", value=1.2, step=0.05, format="%.2f")
    f_excl_lo = z_lo * f_op
    f_excl_hi = z_hi * f_op
    st.caption(f"= {f_excl_lo:.2f} – {f_excl_hi:.2f} Hz")

    st.divider()
    modo_cond = st.selectbox("Condición a graficar",
                             ["Operación", "Transitorio", "Ambas"], index=0)

    st.divider()
    with st.expander("⚙️ Avanzado — Unidades y criterios"):
        F_REF_N = st.number_input("Fuerza de referencia (N por unidad de carga)",
                                  value=9810.0, step=10.0, format="%.0f",
                                  help="SAP2000: 1 Ton = 9810 N. La FRF se normaliza por esta fuerza.")
        factor_despl = st.number_input("Factor de unidad de desplazamiento", value=10.0, step=1.0,
                                       format="%.2f", help="Convierte la salida de SAP a mm. cm→mm: ×10.")
        st.markdown("**Límites ISO 20816-3 (v_RMS, mm/s)**")
        ci1, ci2, ci3 = st.columns(3)
        with ci1: iso_a = st.number_input("A/B", value=1.12, step=0.01, format="%.2f")
        with ci2: iso_b = st.number_input("B/C", value=2.80, step=0.01, format="%.2f")
        with ci3: iso_c = st.number_input("C/D", value=7.10, step=0.01, format="%.2f")

# ══════════════════════════════════════════════════════════════════════════════
# HEADER PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
st.title("📊 Analizador Dinámico — Plataformas de Soporte para Equipos Rotativos")
st.markdown(
    f"**Operación:** {f_op_rpm:.0f} RPM ({f_op:.2f} Hz) | "
    f"**Transitorio f_rd:** {f_rd:.2f} Hz | "
    f"**Zona de exclusión ACI 351.3R:** {f_excl_lo:.2f} – {f_excl_hi:.2f} Hz | "
    f"**K_din aislador:** {K_din:.0f} N/mm"
)

with st.expander("❓ Glosario / FAQ — ¿Qué significa cada variable?"):
    st.markdown(
        "#### Frecuencias\n"
        "| Variable | Significado |\n"
        "|---|---|\n"
        "| **f_op** | Frecuencia de **operación** del equipo (= RPM/60). Es la frecuencia de giro nominal. |\n"
        "| **f_op (RPM)** | Velocidad de giro nominal en revoluciones por minuto. |\n"
        "| **f_rd** | Frecuencia del **modo de cuerpo rígido** del equipo sobre el aislador: f_rd = (1/2π)·√(K_din/m). Es la frecuencia que el equipo **atraviesa al partir o detenerse** (condición transitoria). En modo *Manual* la otorga el fabricante del aislador. |\n"
        "| **Zona de exclusión** | Banda [0.8·f_op, 1.2·f_op] (ACI 351.3R) donde **no** debe haber frecuencias naturales, para evitar resonancia en operación. |\n"
        "| **f_peor** | Frecuencia dentro de la ventana del transitorio donde el producto FRF·F es máximo (peor caso de partida/parada). |\n"
        "\n"
        "#### Respuesta dinámica\n"
        "| Variable | Significado |\n"
        "|---|---|\n"
        "| **FRF** | *Frequency Response Function*: desplazamiento por unidad de fuerza (mm/Ton), obtenida del Steady-State de SAP2000. |\n"
        "| **FRF_op** | Valor de la FRF evaluado en **f_op** (operación). |\n"
        "| **FRF_peak** | Valor de la FRF en el **peak** dentro de la zona de exclusión. |\n"
        "| **φ (fase)** | Ángulo de fase de la respuesta. |φ|≈90° indica condición de **resonancia**. |\n"
        "| **A** | **Amplitud** de desplazamiento peak: A = FRF · F / F_ref (mm). |\n"
        "| **v_peak** | Velocidad peak: v = 2π·f·A (mm/s). |\n"
        "| **v_RMS** | Velocidad RMS = v_peak/√2. Es la que clasifica ISO 20816-3. |\n"
        "\n"
        "#### Fuerzas y aislador\n"
        "| Variable | Significado |\n"
        "|---|---|\n"
        "| **F_rd** | Fuerza dinámica en la condición transitoria. En *Valor fijo* es **constante** (valor práctico); en *Curva desbalance* varía con f². |\n"
        "| **F_op** | Fuerza dinámica en operación (a f_op). |\n"
        "| **U = m·e** | Desbalance del **rotor** (g·mm): m = masa rotante, e = excentricidad. Genera F(f)=m·e·ω². |\n"
        "| **G (ISO 1940-1)** | Grado de calidad de balanceo (mm/s). Define U = m·1000·G/ω. |\n"
        "| **F_ref** | Fuerza de referencia con que se normalizó la FRF (SAP: 1 Ton = 9810 N). |\n"
        "| **K_din** | Rigidez **dinámica** del aislador (N/mm). |\n"
        "| **K_est** | Rigidez **estática** del sistema en operación: K_est = F_ref / FRF_op. |\n"
        "| **RF = K_est/K_din** | Razón de rigidez (Hutchinson). Debe ser ≥ 10 para aislamiento efectivo. |\n"
        "| **W_pata** | Peso estático por apoyo (kgf = masa en kg). |\n"
        "| **N° de apoyos** | Número de patas; reparte la fuerza total de desbalance por apoyo. |\n"
        "\n"
        "#### Clasificación de amplitudes\n"
        "- **Richart (Fig. 10-1)** y **Blake (Fig. 10-2)**: cartas amplitud–frecuencia para clasificar severidad.\n"
        "- **ISO 20816-3**: límites de v_RMS por zonas A (aceptable) / B (normal) / C (alarma) / D (detener).\n"
        "- **Condiciones evaluadas:** ① **Operación** (en f_op) · ② **Transitorio** (partida/parada) · ③ **Peak en zona de exclusión**."
    )

st.divider()

# Reconstruir zonas ISO con los límites configurables del sidebar
ISO_ZONES = [
    (iso_a, "Zona A — Aceptable",        "#1a9850"),
    (iso_b, "Zona B — Normal",           "#66bd63"),
    (iso_c, "Zona C — Alarma",           "#fee08b"),
    (1e9,   "Zona D — Detener equipo",   "#f46d43"),
]

# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════
tab_data, tab_frf, tab_fase, tab_corto, tab_class, tab_dist, tab_report = st.tabs([
    "📥 Datos SAP2000",
    "📈 FRF — Amplitud",
    "🔄 Análisis de Fase",
    "🔗 Cortocircuito Vibratorio",
    "🗂️ Clasificación Richart / Blake / ISO",
    "🔀 Distorsión Angular",
    "📄 Reporte"
])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1: DATOS DE ENTRADA
# ══════════════════════════════════════════════════════════════════════════════
with tab_data:
    st.header("Datos de entrada — SAP2000 Steady-State")

    col_up, col_ex = st.columns([2, 1])
    with col_up:
        uploaded = st.file_uploader(
            "Cargar archivo Excel exportado de SAP2000 (.xlsx)",
            type=["xlsx", "csv"],
            help="Exportar desde SAP2000: Display → Show Tables → Analysis Results → Joint Displacements"
        )
    with col_ex:
        st.markdown("**Formato esperado de columnas:**")
        st.code("Joint | OutputCase | CaseType | StepType | Freq | U1 | U2 | U3", language="")
        st.caption("StepTypes: Mag at Freq / Real at Freq / Imag at Freq")

    with st.expander("O ingresar datos directamente (Formato CSV):"):
        sample_csv = """Joint,OutputCase,CaseType,StepType,Freq,U1,U2,U3,R1,R2,R3
C11,STST_X,LinSteadyState,Mag at Freq,48.33,2.514E-02,1.2E-03,8.5E-04,0,0,0
C11,STST_X,LinSteadyState,Real at Freq,48.33,2.480E-02,1.1E-03,8.0E-04,0,0,0
C11,STST_X,LinSteadyState,Imag at Freq,48.33,-4.2E-03,-2.1E-04,-2.0E-04,0,0,0
C14,STST_X,LinSteadyState,Mag at Freq,57.30,9.787E-02,3.1E-03,9.2E-04,0,0,0
C14,STST_X,LinSteadyState,Real at Freq,57.30,5.9E-03,2.8E-04,8.5E-05,0,0,0
C14,STST_X,LinSteadyState,Imag at Freq,57.30,-9.784E-02,-3.1E-03,-9.2E-04,0,0,0"""

        manual_data = st.text_area(
            "Pegar datos CSV (Joint, OutputCase, CaseType, StepType, Freq, U1, U2, U3, ...)",
            value=sample_csv, height=200
        )

    @st.cache_data
    def load_data(file_bytes, file_name, manual_text):
        try:
            if file_bytes:
                if file_name.endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(file_bytes))
                else:
                    df = pd.read_excel(io.BytesIO(file_bytes), skiprows=2)
                    df.columns = ['Joint','OutputCase','CaseType','StepType',
                                  'Freq','U1','U2','U3','R1','R2','R3']
            else:
                df = pd.read_csv(io.StringIO(manual_text))
            for col in ['Freq','U1','U2','U3']:
                df[col] = pd.to_numeric(df[col], errors='coerce')
            return df.dropna(subset=['Freq'])
        except Exception as e:
            st.error(f"Error al cargar datos: {e}")
            return pd.DataFrame()

    file_bytes = uploaded.read() if uploaded else None
    file_name  = uploaded.name  if uploaded else ""
    df_raw = load_data(file_bytes, file_name, manual_data)
    dir_override = {}

    if not df_raw.empty:
        st.success(f"✅ {len(df_raw)} filas cargadas | "
                   f"Joints: {sorted(df_raw['Joint'].unique())} | "
                   f"Casos: {list(df_raw['OutputCase'].unique())} | "
                   f"StepTypes: {list(df_raw['StepType'].unique())}")

        col_fil1, col_fil2 = st.columns(2)
        with col_fil1:
            casos_disp = sorted(df_raw['OutputCase'].unique())
            casos_sel  = st.multiselect("Casos a analizar", casos_disp, default=casos_disp)
        with col_fil2:
            joints_disp = sorted(df_raw['Joint'].unique())
            joints_sel  = st.multiselect("Joints a analizar", joints_disp, default=joints_disp)

        df_raw = df_raw[df_raw['OutputCase'].isin(casos_sel) &
                        df_raw['Joint'].isin(joints_sel)]
        st.dataframe(df_raw.head(30), use_container_width=True, height=250, hide_index=True)

        with st.expander("🧭 Mapeo de dirección por caso (override)"):
            st.caption("Columna de respuesta (U1/U2/U3) y etiqueta de dirección (X/Y/Z) por caso. "
                       "Por defecto se detecta del nombre del caso; Z = vertical (usa F vertical).")
            dir_opts = ['U1', 'U2', 'U3']; lbl_opts = ['X', 'Y', 'Z']
            for caso in casos_sel:
                ac_col, ac_dir = auto_dir(caso)
                cc1, cc2, cc3 = st.columns([2, 1, 1])
                with cc1: st.markdown(f"**{caso}**")
                with cc2:
                    col_sel = st.selectbox("Columna", dir_opts,
                                           index=dir_opts.index(ac_col),
                                           key=f"colmap_{caso}", label_visibility="collapsed")
                with cc3:
                    di = lbl_opts.index(ac_dir) if ac_dir in lbl_opts else 0
                    dir_sel = st.selectbox("Dir", lbl_opts, index=di,
                                           key=f"dirmap_{caso}", label_visibility="collapsed")
                dir_override[caso] = (col_sel, dir_sel)
    else:
        st.warning("Cargue datos para continuar.")
        casos_sel = []; joints_sel = []; df_raw = pd.DataFrame()

# ══════════════════════════════════════════════════════════════════════════════
# PROCESAMIENTO CENTRAL
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def procesar(df_raw_json, f_op_, f_excl_lo_, f_excl_hi_, f_rd_, factor_despl_, dir_override_json):
    if not df_raw_json:
        return {}, pd.DataFrame()
    df = pd.read_json(io.StringIO(df_raw_json))
    mag_df  = df[df['StepType']=='Mag at Freq'].copy()
    real_df = df[df['StepType']=='Real at Freq'].copy()
    imag_df = df[df['StepType']=='Imag at Freq'].copy()

    # Mapeo case → (col_resp, dir_label): override del usuario o detección automática
    dir_ovr = json.loads(dir_override_json) if dir_override_json else {}
    case_dir = {}
    for caso in df['OutputCase'].unique():
        if str(caso) in dir_ovr:
            case_dir[caso] = tuple(dir_ovr[str(caso)])
        else:
            case_dir[caso] = auto_dir(caso)

    resultados = {}  # {(caso, joint): {...}}

    for caso in df['OutputCase'].unique():
        col_resp, dir_lbl = case_dir.get(caso, ('U1', caso))
        for joint in df['Joint'].unique():
            mj = mag_df[(mag_df['OutputCase']==caso)&(mag_df['Joint']==joint)].sort_values('Freq')
            rj = real_df[(real_df['OutputCase']==caso)&(real_df['Joint']==joint)].sort_values('Freq')
            ij = imag_df[(imag_df['OutputCase']==caso)&(imag_df['Joint']==joint)].sort_values('Freq')
            if mj.empty: continue

            freqs   = mj['Freq'].values
            frf_mm  = np.abs(mj[col_resp].values) * factor_despl_  # → mm, /unidad de carga
            fase_deg= np.degrees(np.arctan2(
                ij[col_resp].values if not ij.empty else np.zeros_like(freqs),
                rj[col_resp].values if not rj.empty else np.zeros_like(freqs)
            ))

            # Peak de OPERACIÓN: se busca el máximo de |H(f)| DENTRO de la zona de
            # exclusión [0.8·fop, 1.2·fop]. La frecuencia de operación es un único
            # valor; lo que importa es si una resonancia estructural cae en esa franja.
            mask_zona = (freqs >= f_excl_lo_) & (freqs <= f_excl_hi_)
            zona_con_datos = bool(mask_zona.any())
            mask = mask_zona if zona_con_datos else np.ones(len(freqs), bool)

            idx_pk  = np.argmax(frf_mm[mask])
            f_pk    = freqs[mask][idx_pk]
            frf_pk  = frf_mm[mask][idx_pk]
            fase_pk = fase_deg[mask][idx_pk]

            idx_op  = np.argmin(np.abs(freqs - f_op_))
            frf_op  = frf_mm[idx_op]
            fase_op = fase_deg[idx_op]

            # Transitorio (partida/parada): FRF en el cruce resonante del aislador f_rd
            idx_rd  = np.argmin(np.abs(freqs - f_rd_))
            frf_rd  = frf_mm[idx_rd]
            fase_rd = fase_deg[idx_rd]
            f_rd_en_rango = bool(freqs.min() <= f_rd_ <= freqs.max())

            # El peak ya está en la zona de exclusión (por construcción de la máscara).
            # La resonancia se confirma con el ángulo de fase ≈ 90° y amplitud relevante.
            en_zona = zona_con_datos
            dist_90 = abs(abs(fase_pk) - 90)

            if not zona_con_datos:
                diagnostico = "Sin datos en zona excl."
            elif dist_90 < 15 and frf_pk > 0.05:
                diagnostico = "Resonancia confirmada"
            elif dist_90 < 30 and frf_pk > 0.05:
                diagnostico = "Probable"
            else:
                diagnostico = "Sin resonancia clara"

            resultados[(caso, joint)] = {
                'caso': caso, 'joint': joint, 'dir': dir_lbl,
                'freqs': freqs.tolist(), 'frf_mm': frf_mm.tolist(),
                'fase_deg': fase_deg.tolist(),
                'f_pk': f_pk, 'frf_pk': frf_pk, 'fase_pk': fase_pk,
                'frf_op': frf_op, 'fase_op': fase_op,
                'f_rd': f_rd_, 'frf_rd': frf_rd, 'fase_rd': fase_rd,
                'f_rd_en_rango': f_rd_en_rango,
                'en_zona': en_zona, 'dist_90': dist_90,
                'diagnostico': diagnostico,
            }

    # DataFrame resumen
    rows = []
    for (caso, joint), r in resultados.items():
        rows.append({
            'Caso': caso, 'Joint': joint, 'Dir': r['dir'],
            'f_peak (Hz)': round(r['f_pk'], 2),
            'FRF_peak (mm/T)': round(r['frf_pk'], 4),
            'FRF_op (mm/T)': round(r['frf_op'], 4),
            'Fase_peak (°)': round(r['fase_pk'], 1),
            'En zona excl.': '✓ Sí' if r['en_zona'] else 'No',
            'Diagnóstico': r['diagnostico'],
        })
    return resultados, pd.DataFrame(rows)

if not df_raw.empty:
    resultados, df_res = procesar(
        df_raw.to_json(), f_op, f_excl_lo, f_excl_hi, f_rd, factor_despl,
        json.dumps({str(k): list(v) for k, v in dir_override.items()}, sort_keys=True)
    )
else:
    resultados, df_res = {}, pd.DataFrame()

# ── Colores por joint ─────────────────────────────────────────────────────────
PALETA = plt.cm.tab20.colors
def color_joint(joints_list, joint):
    idx = sorted(set(joints_list)).index(joint) % len(PALETA)
    return PALETA[idx]

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2: FRF — AMPLITUD
# ══════════════════════════════════════════════════════════════════════════════
with tab_frf:
    st.header("FRF — Amplitud |U| por caso y Joint")
    if not resultados:
        st.info("Cargue datos en la pestaña 'Datos SAP2000' para ver los gráficos.")
    else:
        casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
        joints_uniq = sorted(set(r['joint'] for r in resultados.values()))

        cfg1, cfg2 = st.columns(2)
        with cfg1:
            n_cols = st.slider("Columnas por fila", 1, 3, min(len(casos_uniq), 3))
        with cfg2:
            alto_fila_frf = st.slider("Alto por fila (px)", 250, 800, 380, 10,
                                      help="Controla la altura vertical de cada fila de "
                                           "subgráficos FRF — Amplitud.")
        fig_frf = make_subplots(
            rows=math.ceil(len(casos_uniq)/n_cols), cols=n_cols,
            subplot_titles=casos_uniq, shared_yaxes=False
        )

        for ci, caso in enumerate(casos_uniq):
            row_p = ci // n_cols + 1; col_p = ci % n_cols + 1
            grupo = {k: v for k, v in resultados.items() if k[0] == caso}
            for (c, joint), r in grupo.items():
                col_hex = '#%02x%02x%02x' % tuple(
                    int(x*255) for x in color_joint(joints_uniq, joint)[:3])
                fig_frf.add_trace(go.Scatter(
                    x=r['freqs'], y=r['frf_mm'],
                    name=joint, legendgroup=joint,
                    showlegend=(ci == 0),
                    line=dict(color=col_hex, width=2),
                    mode='lines',
                    hovertemplate=f"<b>{joint}</b><br>f=%{{x:.2f}} Hz<br>FRF=%{{y:.4f}} mm/T<extra></extra>"
                ), row=row_p, col=col_p)
                # Marcar peak
                fig_frf.add_trace(go.Scatter(
                    x=[r['f_pk']], y=[r['frf_pk']],
                    mode='markers', showlegend=False,
                    marker=dict(color=col_hex, size=8, symbol='circle'),
                    hovertemplate=f"<b>{joint} peak</b><br>f={r['f_pk']:.1f} Hz<br>{r['frf_pk']:.4f} mm/T<extra></extra>"
                ), row=row_p, col=col_p)

            # Zona de exclusión
            fig_frf.add_vrect(
                x0=f_excl_lo, x1=f_excl_hi,
                fillcolor="orange", opacity=0.12, line_width=0,
                row=row_p, col=col_p
            )
            fig_frf.add_vline(
                x=f_op, line_dash="dash", line_color="red", opacity=0.7,
                row=row_p, col=col_p
            )

        fig_frf.update_layout(
            height=alto_fila_frf*math.ceil(len(casos_uniq)/n_cols),
            title_text="FRF |U| — Desplazamiento por unidad de fuerza (mm/Ton)",
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            font=dict(family="Arial"),
            hovermode="x unified"
        )
        fig_frf.update_xaxes(title_text="Frecuencia (Hz)", range=[
            df_raw['Freq'].min()*0.95 if not df_raw.empty else 30,
            df_raw['Freq'].max()*1.05 if not df_raw.empty else 80
        ])
        fig_frf.update_yaxes(title_text="|U| (mm/Ton)", rangemode="tozero")
        st.plotly_chart(fig_frf, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3: ANÁLISIS DE FASE
# ══════════════════════════════════════════════════════════════════════════════
with tab_fase:
    st.header("Ángulo de fase φ — Criterio de resonancia (|φ| ≈ 90°)")
    if not resultados:
        st.info("Cargue datos para ver el análisis de fase.")
    else:
        casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
        joints_uniq = sorted(set(r['joint'] for r in resultados.values()))
        n_cols = min(len(casos_uniq), 3)

        fig_fase = make_subplots(
            rows=math.ceil(len(casos_uniq)/n_cols), cols=n_cols,
            subplot_titles=casos_uniq
        )
        for ci, caso in enumerate(casos_uniq):
            row_p = ci // n_cols + 1; col_p = ci % n_cols + 1
            grupo = {k: v for k, v in resultados.items() if k[0] == caso}
            for (c, joint), r in grupo.items():
                col_hex = '#%02x%02x%02x' % tuple(
                    int(x*255) for x in color_joint(joints_uniq, joint)[:3])
                fig_fase.add_trace(go.Scatter(
                    x=r['freqs'], y=r['fase_deg'],
                    name=joint, legendgroup=joint,
                    showlegend=(ci == 0),
                    line=dict(color=col_hex, width=2),
                    mode='lines',
                    hovertemplate=f"<b>{joint}</b><br>f=%{{x:.2f}} Hz<br>φ=%{{y:.1f}}°<extra></extra>"
                ), row=row_p, col=col_p)
                fig_fase.add_trace(go.Scatter(
                    x=[r['f_pk']], y=[r['fase_pk']],
                    mode='markers+text',
                    text=[f"{r['fase_pk']:.0f}°"],
                    textposition="top right",
                    showlegend=False,
                    marker=dict(color=col_hex, size=8),
                    textfont=dict(size=9, color=col_hex)
                ), row=row_p, col=col_p)

            fig_fase.add_vrect(
                x0=f_excl_lo, x1=f_excl_hi,
                fillcolor="orange", opacity=0.12, line_width=0,
                row=row_p, col=col_p
            )
            fig_fase.add_vline(x=f_op, line_dash="dash", line_color="red", opacity=0.7, row=row_p, col=col_p)
            for y_ref in [-90, 90]:
                fig_fase.add_hline(y=y_ref, line_dash="dot", line_color="purple",
                                   opacity=0.6, row=row_p, col=col_p)

        fig_fase.update_layout(
            height=380*math.ceil(len(casos_uniq)/n_cols),
            title_text="Ángulo de fase φ — Líneas ±90° = criterio de resonancia",
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            font=dict(family="Arial"),
        )
        fig_fase.update_yaxes(range=[-220, 220], title_text="φ (°)")
        fig_fase.update_xaxes(title_text="Frecuencia (Hz)")
        st.plotly_chart(fig_fase, use_container_width=True)

        # Tabla diagnóstico
        st.subheader("Resumen de Diagnóstico de Resonancia")
        if not df_res.empty:
            def color_diag(val):
                if "confirmada" in str(val): return "background-color:#FEE2E2; color:#C0392B; font-weight:bold"
                if "Probable"   in str(val): return "background-color:#FEF3C7; color:#D97706"
                if "Posible"    in str(val): return "background-color:#FFFBEB; color:#92400E"
                return "background-color:#F0FDF4; color:#16A34A"
            styled = df_res.style.map(color_diag, subset=['Diagnóstico'])
            st.dataframe(styled, use_container_width=True, hide_index=True)

# ══════════════════════════════════════════════════════════════════════════════
# TAB: CORTOCIRCUITO VIBRATORIO
# ══════════════════════════════════════════════════════════════════════════════
with tab_corto:
    st.header("Cortocircuito Vibratorio — Rigidez relativa (ACI 351.3R-04)")
    if not resultados:
        st.info("Cargue datos para evaluar el cortocircuito vibratorio.")
    else:
        st.markdown(
            "Criterio **10:1** (ACI 351.3R-04, §3.3): la rigidez dinámica de la estructura debe ser "
            "≥ 10× la del aislador para que éste desacople la máquina. K_est se obtiene de la FRF en "
            "operación:  **K_est = F_ref / FRF_op**. El caso gobernante es el de mayor FRF_op (menor rigidez)."
        )
        ratio_obj = st.number_input("Ratio objetivo (criterio)", value=10.0, step=1.0, format="%.0f")

        rows_cc = []
        for (caso, joint), r in resultados.items():
            frf_op = r['frf_op']
            if frf_op <= 0:
                continue
            k_est = F_REF_N / frf_op
            rows_cc.append({
                'Caso': caso, 'Joint': joint, 'Dir': r['dir'],
                'FRF_op (mm/T)': round(frf_op, 4),
                'f_pk (Hz)': round(r['f_pk'], 1),
                'K_est (N/mm)': round(k_est, 0),
                'K_est/K_din': round(k_est / K_din, 2),
                'Cumple 10:1': '✓ Sí' if k_est / K_din >= ratio_obj else '✗ No',
            })
        df_cc = pd.DataFrame(rows_cc)
        if df_cc.empty:
            st.warning("No hay FRF en operación válida (>0) para calcular K_est.")
        else:
            gov = df_cc.loc[df_cc['K_est/K_din'].idxmin()]
            ratio_min = gov['K_est/K_din']
            cumple = ratio_min >= ratio_obj

            m1, m2 = st.columns(2)
            m1.metric("Ratio mínimo K_est/K_din", f"{ratio_min:.2f}")
            m2.metric("K_din aislador", f"{K_din:.0f} N/mm")
            if cumple:
                st.success(f"✓ Cumple: K_est/K_din = {ratio_min:.2f} ≥ {ratio_obj:.0f}.")
            else:
                st.error(
                    f"✗ No cumple (cortocircuito): K_est/K_din = {ratio_min:.2f} < {ratio_obj:.0f}, "
                    f"gobernado por {gov['Caso']} / {gov['Joint']} (dir {gov['Dir']}). La estructura "
                    f"tiene solo el {ratio_min/ratio_obj*100:.0f}% de la rigidez mínima; el aislador "
                    f"no logra desacoplar la máquina de la estructura."
                )

            def color_cc(v):
                return ("background-color:#FEE2E2;color:#C0392B;font-weight:bold"
                        if "No" in str(v) else "background-color:#F0FDF4;color:#166534")
            st.dataframe(df_cc.style.map(color_cc, subset=['Cumple 10:1']),
                         use_container_width=True, hide_index=True, height=320)

            st.subheader("Transmisibilidad en el cruce de resonancia")
            st.caption("Fracción de fuerza que cruza el aislador hacia la estructura: "
                       "T = √(1+(2ξr)²) / √((1−r²)²+(2ξr)²), con r = f_op/f_res.")
            t1, t2 = st.columns(2)
            with t1:
                xi = st.number_input("Amortiguamiento ξ (%)", value=3.0, step=0.5, format="%.1f") / 100
            with t2:
                f_res = st.number_input("f_res (Hz)", value=float(gov['f_pk (Hz)']), step=0.1, format="%.1f")
            rr = f_op / f_res if f_res > 0 else 0.0
            T = math.sqrt(1 + (2*xi*rr)**2) / math.sqrt((1 - rr**2)**2 + (2*xi*rr)**2)
            st.metric("Transmisibilidad T", f"{T:.2f}")
            st.caption(f"r = f_op/f_res = {f_op:.2f}/{f_res:.1f} = {rr:.3f}  →  ~{min(T,9.99)*100:.0f}% de la "
                       f"fuerza se transmite a la estructura.")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 4: CLASIFICACIÓN RICHART / BLAKE / ISO
# ══════════════════════════════════════════════════════════════════════════════
with tab_class:
    st.header("Clasificación de amplitudes — Richart, Blake e ISO 20816-3")
    if not resultados:
        st.info("Cargue datos para ver la clasificación.")
    else:
        st.markdown("**Configurar fuerzas reales para calcular amplitudes:**")
        col_a, col_b, col_c, col_d = st.columns(4)
        with col_a: F_op_V_c  = st.number_input("F op V (N)", value=F_op_V,  key="fopv_c")
        with col_b: F_op_H_c  = st.number_input("F op H (N)", value=F_op_H,  key="foph_c")
        with col_c: F_rd_V_c  = st.number_input("F rd V (N)", value=F_rd_V,  key="frdv_c")
        with col_d: F_rd_H_c  = st.number_input("F rd H (N)", value=F_rd_H,  key="frdh_c")

        # Tres condiciones → una tabla cada una (cada una con Richart + Blake + ISO)
        rows_op, rows_tr, rows_pk = [], [], []
        for (caso, joint), r in resultados.items():
            dir_lbl = r['dir']
            F_op_cal = F_op_V_c if dir_lbl == 'Z' else F_op_H_c
            F_rd_cal = F_rd_V_c if dir_lbl == 'Z' else F_rd_H_c

            # ── ① Operación: FRF a f_op × F_op ──
            u_op  = r['frf_op'] * F_op_cal / F_REF_N
            vp_op = v_peak(u_op, f_op); vr_op = v_rms(u_op, f_op)
            rows_op.append({
                'Caso': caso, 'Joint': joint, 'Dir': dir_lbl,
                'f_op (Hz)': round(f_op, 2),
                'FRF_op (mm/T)': round(r['frf_op'], 5),
                'F_op (N)': round(F_op_cal, 0),
                'A_op (mm)': round(u_op, 5),
                'v_op (mm/s)': round(vp_op, 3),
                'vRMS_op (mm/s)': round(vr_op, 3),
                'Richart': classify(vp_op, RICHART_ZONES)[0],
                'Blake':   classify(vp_op, BLAKE_ZONES)[0],
                'ISO':     classify(vr_op, ISO_ZONES)[0],
            })

            # ── ② Transitorio: peor caso del tramo (+ punto f_rd) ──
            u_rd, f_rd_eff, frf_rd_used, F_rd_used = transitorio_uf(
                r, F_rd_cal, modo_fuerza_tr, U_gmm, F_REF_N, f_tr_lo, f_tr_hi, n_apoyos)
            u_frd, f_frd, F_frd, frd_en_rango = transitorio_en_frd(
                r, F_rd_cal, modo_fuerza_tr, U_gmm, F_REF_N, n_apoyos)
            vp_rd = v_peak(u_rd, f_rd_eff); vr_rd = v_rms(u_rd, f_rd_eff)
            rows_tr.append({
                'Caso': caso, 'Joint': joint, 'Dir': dir_lbl,
                # Peor caso del tramo (la fuerza F_peor se evalúa en f_peor)
                'f_peor (Hz)': round(f_rd_eff, 2),
                'FRF_peor (mm/T)': round(frf_rd_used, 5),
                'F_peor (N)': round(F_rd_used, 0),
                'A_peor (mm)': round(u_rd, 4),
                # Punto de cruce del aislador f_rd (la fuerza F@f_rd se evalúa en f_rd:
                # en modo desbalance F∝f², por lo que aquí es mucho menor que F_peor)
                'F@f_rd (N)': round(F_frd, 0),
                'A@f_rd (mm)': round(u_frd, 4) if frd_en_rango else float('nan'),
                'v_rd (mm/s)': round(vp_rd, 2),
                'vRMS_rd (mm/s)': round(vr_rd, 2),
                'Richart': classify(vp_rd, RICHART_ZONES)[0],
                'Blake':   classify(vp_rd, BLAKE_ZONES)[0],
                'ISO':     classify(vr_rd, ISO_ZONES)[0],
            })

            # ── ③ Peak en zona de exclusión: FRF_pk × F_op, evaluado a f_pk ──
            u_pk  = r['frf_pk'] * F_op_cal / F_REF_N
            vp_pk = v_peak(u_pk, r['f_pk']); vr_pk = v_rms(u_pk, r['f_pk'])
            rows_pk.append({
                'Caso': caso, 'Joint': joint, 'Dir': dir_lbl,
                'f_peak (Hz)': round(r['f_pk'], 2),
                'FRF_peak (mm/T)': round(r['frf_pk'], 5),
                'Fase_peak (°)': round(r['fase_pk'], 1),
                'F_op (N)': round(F_op_cal, 0),
                'A_peak (mm)': round(u_pk, 5),
                'v_peak (mm/s)': round(vp_pk, 3),
                'vRMS_peak (mm/s)': round(vr_pk, 3),
                'Richart': classify(vp_pk, RICHART_ZONES)[0],
                'Blake':   classify(vp_pk, BLAKE_ZONES)[0],
                'ISO':     classify(vr_pk, ISO_ZONES)[0],
                'En zona': '✓' if r['en_zona'] else '—',
                'Diagnóstico': r['diagnostico'],
            })

        df_op = pd.DataFrame(rows_op)
        df_tr = pd.DataFrame(rows_tr)
        df_pk = pd.DataFrame(rows_pk)

        # Advertencia: la ventana del transitorio no queda cubierta por los datos cargados
        if resultados:
            f_min_d = df_raw['Freq'].min(); f_max_d = df_raw['Freq'].max()
            win_lo, win_hi = min(f_tr_lo, f_tr_hi), max(f_tr_lo, f_tr_hi)
            sin_solape = win_hi < f_min_d or win_lo > f_max_d
            f_rd_fuera = not (f_min_d <= f_rd <= f_max_d)
            if sin_solape:
                f_borde = f_min_d if win_hi < f_min_d else f_max_d
                st.warning(
                    f"⚠️ La ventana del transitorio ({win_lo:.2f}–{win_hi:.2f} Hz) está **fuera "
                    f"del rango de datos cargados** ({f_min_d:.1f}–{f_max_d:.1f} Hz). Se evalúa el "
                    f"extremo más cercano disponible: **{f_borde:.1f} Hz**, por lo que los "
                    f"resultados de transitorio **no son válidos**. Sube un barrido SAP2000 que "
                    f"incluya esa banda (idealmente 0 → ~50 Hz)."
                )
            elif f_rd_fuera or win_lo < f_min_d:
                st.warning(
                    f"⚠️ Parte de la ventana del transitorio ({win_lo:.2f}–{win_hi:.2f} Hz) cae "
                    f"**bajo el mínimo de datos cargados** ({f_min_d:.1f} Hz) — el cruce del "
                    f"aislador f_rd = {f_rd:.2f} Hz {'no está' if f_rd_fuera else 'está'} dentro de "
                    f"los datos. El peak transitorio se busca solo en el tramo con datos "
                    f"({max(win_lo, f_min_d):.1f}–{min(win_hi, f_max_d):.1f} Hz). Para capturar el "
                    f"cruce del aislador, sube un barrido que incluya la baja frecuencia."
                )

        # Mostrar con colores
        crit_cols = ['Richart', 'Blake', 'ISO']
        def color_class(val):
            # Fija SIEMPRE color de fondo y de texto (oscuro) para asegurar contraste
            # tanto en tema claro como oscuro de Streamlit.
            val = str(val)
            if any(x in val for x in ["Peligro","Zona D","Severa","E — "]):
                return "background-color:#FEE2E2;color:#991B1B;font-weight:bold"
            if any(x in val for x in ["Precaución","Límite","Zona C","C — ","D — ","Molesta"]):
                return "background-color:#FEF3C7;color:#92400E;font-weight:bold"
            if any(x in val for x in ["Fácilmente","B — ","Zona B"]):
                return "background-color:#FEF9C3;color:#854D0E"
            if any(x in val for x in ["No perceptible","Zona A","A — ","Apenas"]):
                return "background-color:#DCFCE7;color:#166534"
            return "color:#1F2937"
        
        st.subheader("① Condición de operación (f_op)")
        st.dataframe(df_op.style.map(color_class, subset=crit_cols),
                     use_container_width=True, height=300, hide_index=True)

        st.subheader("② Condición transitorio (partida / parada)")
        if modo_fuerza_tr.startswith("Curva"):
            st.caption("ℹ️ En modo **curva de desbalance**, la fuerza varía con la frecuencia "
                       "(F = m·e·ω² ∝ f²). Por eso `F_peor` (en f_peor) y `F@f_rd` (en el cruce del "
                       "aislador) son **distintas**: la fuerza centrífuga es mínima a baja f y crece "
                       "con f², así que el peor caso tiende a quedar cerca de operación, no en f_rd.")
        else:
            st.caption("ℹ️ En modo **valor fijo**, la fuerza F_rd es **constante** en todo el tramo "
                       "(`F_peor` = `F@f_rd` = F_rd del fabricante).")
        st.dataframe(df_tr.style.map(color_class, subset=crit_cols),
                     use_container_width=True, height=300, hide_index=True)

        st.subheader("③ Condición peaks en zona de exclusión")
        st.caption(f"Peak estructural de |H| dentro de la zona de exclusión "
                   f"[{f_excl_lo:.1f}–{f_excl_hi:.1f} Hz], excitado por la fuerza de operación F_op. "
                   "Evalúa la severidad si una resonancia estructural cae cerca de la operación "
                   "(ver columnas Fase / En zona / Diagnóstico).")
        st.dataframe(df_pk.style.map(color_class, subset=crit_cols),
                     use_container_width=True, height=300, hide_index=True)

        with st.expander("ℹ️ Cómo se calculan las amplitudes (3 condiciones)"):
            st.markdown(
                "**Amplitud** en cada caso: $A = \\mathrm{FRF} \\cdot F / F_{ref}$, con la FRF "
                "[mm/T] leída de SAP2000 y $F_{ref}$ = "
                f"{F_REF_N:.0f} N (la fuerza unitaria con que se normalizó la FRF).\n\n"
                "**Operación** (`_op`): se evalúa **en f_op**. `FRF_op` es la FRF en esa frecuencia "
                "y se multiplica por la **fuerza de operación** `F_op` (vertical si Dir = Z, "
                "horizontal en caso contrario).\n\n"
                "**Transitorio:** se reportan **dos lecturas** del mismo barrido:\n\n"
                "1. **Peor caso del tramo** (`_peor`): se recorre la **ventana de barrido** "
                f"({min(f_tr_lo, f_tr_hi):.2f}–{max(f_tr_lo, f_tr_hi):.2f} Hz) y se toma el máximo "
                "de $A(f)=\\mathrm{FRF}(f)\\cdot F(f)/F_{ref}$. `f_peor` es la frecuencia donde "
                "ocurre, y `FRF_peor`/`F_peor` los valores usados ahí. **Esta es la que clasifica** "
                "(Richart/ISO `_rd`) por ser conservadora.\n"
                "2. **Valor puntual en el cruce del aislador** (`A@f_rd`): la respuesta evaluada "
                f"exactamente en f_rd = {f_rd:.2f} Hz. Es **NaN** si f_rd cae fuera de los datos "
                "cargados (no hay FRF en esa frecuencia).\n\n"
                "Según el **modelo de fuerza transitoria** del sidebar:\n\n"
                "- **Valor fijo (F_rd):** $F(f)=F_{rd}$ **constante** = fuerza dinámica máxima del "
                "fabricante (vertical u horizontal según Dir, ya por apoyo). El peor caso cae en la "
                "**frecuencia de mayor FRF** del tramo.\n"
                "- **Curva desbalance:** $F(f)=m e\\,\\omega^2 / N_{apoyos}$ (desbalance total del "
                "rotor **repartido entre los apoyos**). Crece con $f^2$, así que el peor caso "
                "pondera FRF y fuerza y suele desplazarse hacia arriba.\n\n"
                "**③ Peaks en zona de exclusión** (`_peak`): toma el peak estructural de |H| dentro "
                "de la zona de exclusión [0.8·fop, 1.2·fop] (columna `FRF_peak` a la frecuencia "
                "`f_peak`) y lo excita con la **fuerza de operación** F_op. Responde *qué tan severa "
                "sería la vibración si una resonancia estructural cae cerca de la operación*. Las "
                "columnas `Fase_peak`, `En zona` y `Diagnóstico` indican si es una resonancia real "
                "(fase ≈ 90°, peak dentro de la zona).\n\n"
                "**Cobertura de datos:** si la ventana del transitorio solo solapa parcialmente con "
                "el barrido SAP, el peor caso se busca **dentro del tramo con datos** — no en el "
                "extremo. Solo si la ventana queda totalmente fuera se cae al punto más cercano (ver "
                "aviso amarillo). Todo es **por nodo y caso**: cada fila tiene su propio peor caso."
            )

        # Control de alto para las gráficas de amplitudes
        alto_graf = st.slider("Alto de las gráficas de amplitudes (px)",
                              min_value=400, max_value=1600, value=850, step=50)

        # Puntos (Op ○ / Tr △) para las descargas en alta calidad (matplotlib),
        # idénticos a los del reporte; respetan el selector "Condición a graficar".
        _joints_all = sorted(set(j for (_, j) in resultados.keys()))
        def _hex_joint_app(j):
            return '#%02x%02x%02x' % tuple(int(x*255) for x in color_joint(_joints_all, j)[:3])
        puntos_cl_app = []
        if modo_cond in ["Operación", "Ambas"]:
            for rr in rows_op:
                puntos_cl_app.append((f_op*60, rr['A_op (mm)'], str(rr['Joint']),
                                      _hex_joint_app(rr['Joint']), 'o'))
        if modo_cond in ["Transitorio", "Ambas"]:
            for rr in rows_tr:
                puntos_cl_app.append((rr['f_peor (Hz)']*60, rr['A_peor (mm)'], str(rr['Joint']),
                                      _hex_joint_app(rr['Joint']), '^'))

        def descarga_png_mpl(fig, nombre, etiqueta):
            """Exporta una figura matplotlib (misma calidad que el reporte) como PNG."""
            b = io.BytesIO()
            fig.savefig(b, format='png', dpi=200, bbox_inches='tight')
            b.seek(0); plt.close(fig)
            st.download_button(etiqueta, b, file_name=nombre, mime="image/png", key=nombre)

        # Gráfico de barras v_RMS para ISO
        st.subheader("ISO 20816-3 — v_RMS por nodo y dirección")

        # Ajustar el tope Y dinámicamente según el valor máximo de los datos
        max_vrms = df_op['vRMS_op (mm/s)'].max() if not df_op.empty else 0
        if pd.isna(max_vrms): max_vrms = 0
        iso_top_y = max(max_vrms * 1.2, 10.0) # Asegura al menos ver hasta 10 mm/s
        
        fig_iso = go.Figure()
        ISO_LIMS  = [iso_a, iso_b, iso_c]
        ISO_NAMES = [f"Zona A (≤{iso_a:.2f})", f"Zona B (≤{iso_b:.2f})", f"Zona C (≤{iso_c:.2f})"]
        ISO_COLS  = ["#1a9850",        "#fee08b",         "#f46d43"]
        for y_val, name, col_ in zip(ISO_LIMS, ISO_NAMES, ISO_COLS):
            fig_iso.add_hline(
                y=y_val, line_dash="dash", line_color=col_,
                annotation_text=name, 
                annotation_position="top right",
                annotation_font=dict(color=col_, size=12, family="Arial"),
                line_width=2
            )

        dir_colors = {"X":"#1f77b4","Y":"#e74c3c","Z":"#27ae60"}
        for dir_ in df_op['Dir'].unique():
            df_d = df_op[df_op['Dir']==dir_]
            labels = df_d['Joint'].astype(str) + "<br>(" + df_d['Caso'].astype(str) + ")"
            fig_iso.add_trace(go.Bar(
                name=f"Dir {dir_} — Operación",
                x=labels, y=df_d['vRMS_op (mm/s)'],
                marker_color=dir_colors.get(dir_, "#666"),
                opacity=0.85,
                hovertemplate="<b>%{x}</b><br>v_RMS = %{y:.3f} mm/s<extra></extra>"
            ))

        fig_iso.update_layout(
            title="ISO 20816-3 — Velocidad eficaz v_RMS en operación",
            yaxis_title="v_RMS (mm/s)", xaxis_title="Punto de apoyo",
            yaxis=dict(range=[0, iso_top_y]), # Aplicando el tope dinámico
            barmode="group", height=alto_graf,
            font=dict(family="Arial"),
            legend=dict(orientation="h", yanchor="bottom", y=-0.3)
        )
        fig_iso.add_hrect(y0=0, y1=iso_a, fillcolor="#1a9850", opacity=0.06, line_width=0)
        fig_iso.add_hrect(y0=iso_a, y1=iso_b, fillcolor="#66bd63", opacity=0.06, line_width=0)
        fig_iso.add_hrect(y0=iso_b, y1=iso_c, fillcolor="#fee08b", opacity=0.08, line_width=0)
        fig_iso.add_hrect(y0=iso_c, y1=iso_top_y, fillcolor="#f46d43", opacity=0.08, line_width=0)
        st.plotly_chart(fig_iso, use_container_width=True)
        descarga_png_mpl(fig_iso_barras_mpl(df_op, iso_a, iso_b, iso_c,
                                            "ISO 20816-3 — v_RMS en operación por nodo"),
                         "ISO_20816.png", "⬇️ Descargar ISO (PNG alta calidad)")

        # ══════════════════════════════════════════════════════════════════════
        # FORMATO EXACTO DE EJES LOGARÍTMICOS Y ESTILOS
        # ══════════════════════════════════════════════════════════════════════
        grid_style = dict(showgrid=True, gridwidth=1, gridcolor='rgba(128, 128, 128, 0.25)')
        minor_grid_style = dict(dtick="D1", showgrid=True, gridwidth=0.5, gridcolor='rgba(128, 128, 128, 0.1)')
        
        X_TICK_VALS = [100, 1000, 10000]
        X_TICK_TEXT = ["100", "1000", "10000"]
        Y_TICK_VALS = [0.0001, 0.001, 0.01, 0.1, 1, 10, 100]
        Y_TICK_TEXT = ["0.0001", "0.001", "0.01", "0.1", "1", "10", "100"]

        # ----------------------------------------------------------------------
        # Gráfico Richart Fig 10-1
        # ----------------------------------------------------------------------
        st.subheader("Richart Fig. 10-1 — Puntos de análisis en diagrama log-log")
        fig_rich = go.Figure()

        richart_lines = RICHART_LINES

        for lbl, data in richart_lines.items():
            fig_rich.add_trace(go.Scatter(
                x=data["x"], y=data["y"], mode='lines', name=lbl,
                line=dict(color=data["col"], width=2, dash=data["dash"], shape='linear'),
                showlegend=True, hovertemplate=f"{lbl}<extra></extra>"
            ))

        plotted = set()
        for (caso, joint), r in resultados.items():
            dir_lbl = r['dir']
            F_op_cal = F_op_V_c if dir_lbl == 'Z' else F_op_H_c
            F_rd_cal = F_rd_V_c if dir_lbl == 'Z' else F_rd_H_c
            u_op = r['frf_op'] * F_op_cal / F_REF_N
            u_rd, f_rd_eff, _, _ = transitorio_uf(r, F_rd_cal, modo_fuerza_tr, U_gmm, F_REF_N, f_tr_lo, f_tr_hi, n_apoyos)
            col_ = '#%02x%02x%02x' % tuple(
                int(x*255) for x in color_joint(
                    sorted(set(rv['joint'] for rv in resultados.values())), joint)[:3])

            puntos_a_graficar = []
            if modo_cond in ["Operación", "Ambas"]:
                puntos_a_graficar.append((u_op, f_op, "Op", "circle"))
            if modo_cond in ["Transitorio", "Ambas"]:
                puntos_a_graficar.append((u_rd, f_rd_eff, "Tr", "triangle-up"))

            for u_val, f_val, cond_lbl, sym in puntos_a_graficar:
                key_ = f"{caso}_{joint}_{cond_lbl}"
                if key_ not in plotted and u_val > 0:
                    plotted.add(key_)
                    fig_rich.add_trace(go.Scatter(
                        x=[f_val*60], y=[u_val],
                        mode='markers',
                        name=f"{joint} {dir_lbl} {cond_lbl}",
                        marker=dict(symbol=sym, size=10, color=col_,
                                    line=dict(width=1.5, color="black")),
                        showlegend=False,
                        hovertemplate=(
                            f"<b>{joint} — Dir {dir_lbl} — {cond_lbl}</b><br>"
                            f"f = {f_val:.1f} Hz ({f_val*60:.0f} cpm)<br>"
                            f"A = {u_val:.5f} mm<extra></extra>")
                    ))

        fig_rich.add_vrect(x0=f_excl_lo*60, x1=f_excl_hi*60,
                           fillcolor="orange", opacity=0.15, line_width=0)
        fig_rich.add_vline(x=f_op*60, line_dash="dash", line_color="red", opacity=0.8)

        fig_rich.update_layout(
            title="Richart Fig. 10-1 — Amplitud vs Frecuencia",
            height=alto_graf, font=dict(family="Arial"),
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            margin=dict(l=60, r=40, t=60, b=60), 
            plot_bgcolor='rgba(0,0,0,0)', 
            paper_bgcolor='rgba(0,0,0,0)'
        )
        
        fig_rich.update_xaxes(
            title="Frecuencia (cpm)", type="log", range=[2, 4], 
            tickmode="array", tickvals=X_TICK_VALS, ticktext=X_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        fig_rich.update_yaxes(
            title="Amplitud peak (mm)", type="log", range=[-4.2, 1.2], 
            tickmode="array", tickvals=Y_TICK_VALS, ticktext=Y_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        
        st.plotly_chart(fig_rich, use_container_width=True, config={'toImageButtonOptions': {'format': 'png', 'filename': 'Grafico_Richart', 'scale': 3}})
        descarga_png_mpl(fig_clasif_loglog_mpl(RICHART_LINES, puntos_cl_app, f_excl_lo, f_excl_hi,
                                               f_op, "Richart Fig. 10-1 — Amplitud vs Frecuencia",
                                               (10**-4.2, 10**1.2)),
                         "Richart.png", "⬇️ Descargar Richart (PNG alta calidad)")

        # ----------------------------------------------------------------------
        # Gráfico Blake Fig 10-2
        # ----------------------------------------------------------------------
        st.subheader("Blake Fig. 10-2 — Criterios de severidad vibratoria")
        fig_blake = go.Figure()

        blake_lines = BLAKE_LINES

        for lbl, data in blake_lines.items():
            fig_blake.add_trace(go.Scatter(
                x=data["x"], y=data["y"], mode='lines', name=lbl,
                line=dict(color=data["col"], width=2, dash='solid', shape='linear'),
                showlegend=True, hovertemplate=f"{lbl}<extra></extra>"
            ))

        plotted_blake = set()
        for (caso, joint), r in resultados.items():
            dir_lbl = r['dir']
            F_op_cal = F_op_V_c if dir_lbl == 'Z' else F_op_H_c
            F_rd_cal = F_rd_V_c if dir_lbl == 'Z' else F_rd_H_c
            u_op = r['frf_op'] * F_op_cal / F_REF_N
            u_rd, f_rd_eff, _, _ = transitorio_uf(r, F_rd_cal, modo_fuerza_tr, U_gmm, F_REF_N, f_tr_lo, f_tr_hi, n_apoyos)
            col_ = '#%02x%02x%02x' % tuple(
                int(x*255) for x in color_joint(
                    sorted(set(rv['joint'] for rv in resultados.values())), joint)[:3])

            puntos_a_graficar = []
            if modo_cond in ["Operación", "Ambas"]:
                puntos_a_graficar.append((u_op, f_op, "Op", "circle"))
            if modo_cond in ["Transitorio", "Ambas"]:
                puntos_a_graficar.append((u_rd, f_rd_eff, "Tr", "triangle-up"))

            for u_val, f_val, cond_lbl, sym in puntos_a_graficar:
                key_ = f"{caso}_{joint}_{cond_lbl}"
                if key_ not in plotted_blake and u_val > 0:
                    plotted_blake.add(key_)
                    fig_blake.add_trace(go.Scatter(
                        x=[f_val*60], y=[u_val],
                        mode='markers',
                        name=f"{joint} {dir_lbl} {cond_lbl}",
                        marker=dict(symbol=sym, size=10, color=col_,
                                    line=dict(width=1.5, color="black")),
                        showlegend=False,
                        hovertemplate=(
                            f"<b>{joint} — Dir {dir_lbl} — {cond_lbl}</b><br>"
                            f"f = {f_val:.1f} Hz ({f_val*60:.0f} cpm)<br>"
                            f"A = {u_val:.5f} mm<extra></extra>")
                    ))

        fig_blake.add_vrect(x0=f_excl_lo*60, x1=f_excl_hi*60,
                            fillcolor="orange", opacity=0.15, line_width=0)
        fig_blake.add_vline(x=f_op*60, line_dash="dash", line_color="red", opacity=0.8)
        
        fig_blake.update_layout(
            title="Blake Fig. 10-2 — Amplitud vs Frecuencia",
            height=alto_graf, font=dict(family="Arial"), 
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            margin=dict(l=60, r=40, t=60, b=60), 
            plot_bgcolor='rgba(0,0,0,0)', 
            paper_bgcolor='rgba(0,0,0,0)'
        )
        
        fig_blake.update_xaxes(
            title="Frecuencia (cpm)", type="log", range=[2, 4], 
            tickmode="array", tickvals=X_TICK_VALS, ticktext=X_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        fig_blake.update_yaxes(
            title="Amplitud peak (mm)", type="log", range=[-4.5, 0.6], 
            tickmode="array", tickvals=Y_TICK_VALS, ticktext=Y_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        
        st.plotly_chart(fig_blake, use_container_width=True, config={'toImageButtonOptions': {'format': 'png', 'filename': 'Grafico_Blake', 'scale': 3}})
        descarga_png_mpl(fig_clasif_loglog_mpl(BLAKE_LINES, puntos_cl_app, f_excl_lo, f_excl_hi,
                                               f_op, "Blake Fig. 10-2 — Amplitud vs Frecuencia",
                                               (10**-4.5, 10**0.6)),
                         "Blake.png", "⬇️ Descargar Blake (PNG alta calidad)")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 5: DISTORSIÓN ANGULAR
# ══════════════════════════════════════════════════════════════════════════════
with tab_dist:
    st.header("Distorsión Angular Dinámica — Diagrama Vectorial")
    if not resultados:
        st.info("Cargue datos para ver el análisis de distorsión angular.")
    else:
        casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
        casos_sel_dist = st.multiselect("Seleccionar casos", casos_uniq, default=casos_uniq[:1])
        f_eval = st.number_input("Frecuencia de evaluación (Hz)", value=f_op, step=0.1, format="%.2f")

        if casos_sel_dist:
            F_op_V_d = F_op_V; F_op_H_d = F_op_H

            for caso in casos_sel_dist:
                st.subheader(f"Caso: {caso}")
                grupo = {k: v for k, v in resultados.items() if k[0] == caso}
                if not grupo: continue

                joints_c = sorted(set(k[1] for k in grupo))
                dir_lbl_c = next(iter(grupo.values()))['dir']
                F_f = F_op_V_d if dir_lbl_c == 'Z' else F_op_H_d

                # Obtener fase y amplitud a f_eval
                reales, imags, mags, fases, labels = [], [], [], [], []
                for joint in joints_c:
                    r = grupo.get((caso, joint))
                    if r is None: continue
                    freqs = np.array(r['freqs'])
                    idx   = np.argmin(np.abs(freqs - f_eval))
                    frf   = np.array(r['frf_mm'])
                    fase  = np.array(r['fase_deg'])
                    mag_v = frf[idx] * F_f / F_REF_N
                    phi   = np.radians(fase[idx])
                    reales.append(mag_v * np.cos(phi))
                    imags.append(mag_v * np.sin(phi))
                    mags.append(mag_v)
                    fases.append(fase[idx])
                    labels.append(joint)

                if not labels:
                    st.warning("Sin datos para este caso.")
                    continue

                col_v1, col_v2 = st.columns([2, 1])
                with col_v1:
                    fig_v = go.Figure()
                    colors_v = plt.cm.Set1.colors
                    for k, (re, im, joint) in enumerate(zip(reales, imags, labels)):
                        col_ = '#%02x%02x%02x' % tuple(int(x*255) for x in colors_v[k % len(colors_v)][:3])
                        fig_v.add_trace(go.Scatter(
                            x=[0, re], y=[0, im], mode='lines+markers',
                            name=joint, line=dict(color=col_, width=3),
                            marker=dict(symbol=['circle','arrow'], size=[5,14],
                                        color=col_, angleref='previous'),
                        ))
                        fig_v.add_annotation(
                            x=re*1.2, y=im*1.2, text=f"<b>{joint}</b><br>{fases[k]:.0f}°",
                            showarrow=False, font=dict(size=11, color=col_)
                        )
                    fig_v.add_hline(y=0, line_color="gray", line_width=0.5)
                    fig_v.add_vline(x=0, line_color="gray", line_width=0.5)
                    lim = max(abs(max(reales+imags, key=abs))*1.6, 1e-9)
                    fig_v.update_layout(
                        title=f"Vectores de desplazamiento — {caso} @ {f_eval:.1f} Hz",
                        xaxis=dict(title="Re (mm)", range=[-lim, lim], scaleanchor="y"),
                        yaxis=dict(title="Im (mm)", range=[-lim, lim]),
                        height=450, font=dict(family="Arial"),
                        showlegend=True
                    )
                    st.plotly_chart(fig_v, use_container_width=True)

                with col_v2:
                    st.markdown("**Tabla de vectores**")
                    df_v = pd.DataFrame({
                        'Joint': labels,
                        'Mag (mm)': [round(m, 5) for m in mags],
                        'Fase (°)': [round(f, 1) for f in fases],
                    })
                    st.dataframe(df_v, use_container_width=True, hide_index=True)

                    # Diferencial de fases
                    if len(fases) >= 2:
                        st.markdown("**Diferenciales de fase:**")
                        for i in range(len(labels)-1):
                            df_i = abs(fases[i] - fases[i+1])
                            if df_i > 180: df_i = 360 - df_i
                            emoji = "🔴" if df_i > 90 else "🟡" if df_i > 45 else "🟢"
                            st.markdown(
                                f"{emoji} **{labels[i]}–{labels[i+1]}:** {df_i:.1f}°"
                            )

# ══════════════════════════════════════════════════════════════════════════════
# TAB 6: REPORTE
# ══════════════════════════════════════════════════════════════════════════════
with tab_report:
    st.header("Generación de Reporte")
    if not resultados:
        st.info("Cargue datos para generar el reporte.")
    else:
        col_rp1, col_rp2 = st.columns(2)
        with col_rp1:
            titulo_rep  = st.text_input("Título del reporte", "Auditoría Dinámica — Plataforma de Soporte")
            proyecto    = st.text_input("Proyecto / Referencia", "EST-P730_RA-DISEÑO")
            autor       = st.text_input("Elaborado por", "")
        with col_rp2:
            revisor     = st.text_input("Revisado por", "")
            revision    = st.text_input("Revisión", "A")
            incluir_figs = st.checkbox("Incluir gráficos en el reporte", value=True)

        if st.button("📄 Generar Reporte Word (.docx)", type="primary"):
            with st.spinner("Generando reporte..."):
                # ── Crear documento Word ──────────────────────────────────────
                doc = Document()
                sec = doc.sections[0]
                sec.page_width  = Cm(21)
                sec.page_height = Cm(29.7)
                sec.left_margin = sec.right_margin = Cm(2.5)
                sec.top_margin  = sec.bottom_margin = Cm(2.5)

                def add_h1(text):
                    p = doc.add_paragraph()
                    r = p.add_run(text)
                    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
                    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after  = Pt(6)
                    pf = p.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.5

                def add_para(text):
                    p = doc.add_paragraph()
                    r = p.add_run(text)
                    r.font.name = 'Arial'; r.font.size = Pt(11)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p.paragraph_format.line_spacing = 1.5

                def add_h2(text):
                    p = doc.add_paragraph()
                    r = p.add_run(text)
                    r.font.name='Arial'; r.font.size=Pt(12); r.font.bold=True
                    r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
                    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)

                def add_df_table(df, fs=8):
                    if df is None or df.empty:
                        add_para("(Sin datos para esta sección.)"); return
                    t = doc.add_table(rows=1+len(df), cols=len(df.columns))
                    t.style = 'Table Grid'
                    for ci, col_n in enumerate(df.columns):
                        c = t.rows[0].cells[ci]; c.text = str(col_n)
                        rr = c.paragraphs[0].runs[0]
                        rr.font.bold=True; rr.font.size=Pt(fs); rr.font.name='Arial'
                        rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        tcp = c._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
                        shd.set(q('w:val'),'clear'); shd.set(q('w:color'),'auto'); shd.set(q('w:fill'),'1B2A4A')
                        tcp.append(shd)
                    for ri, row in df.reset_index(drop=True).iterrows():
                        for ci, val in enumerate(row):
                            c = t.rows[ri+1].cells[ci]; c.text = str(val)
                            rr = c.paragraphs[0].runs[0]; rr.font.size=Pt(fs); rr.font.name='Arial'

                # Portada
                doc.add_paragraph()
                p = doc.add_paragraph()
                r = p.add_run(titulo_rep.upper())
                r.font.name='Arial'; r.font.size=Pt(20); r.font.bold=True
                r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

                doc.add_paragraph()
                tabla_port = doc.add_table(rows=6, cols=2)
                tabla_port.style = 'Table Grid'
                datos_port = [
                    ("Proyecto", proyecto),
                    ("Elaborado por", autor),
                    ("Revisado por", revisor),
                    ("Revisión", revision),
                    ("Frecuencia de operación", f"{f_op:.2f} Hz ({f_op_rpm:.0f} RPM)"),
                    ("Zona de exclusión", f"{f_excl_lo:.2f} – {f_excl_hi:.2f} Hz"),
                ]
                for ri, (k, v) in enumerate(datos_port):
                    tabla_port.rows[ri].cells[0].text = k
                    tabla_port.rows[ri].cells[1].text = v
                    tabla_port.rows[ri].cells[0].paragraphs[0].runs[0].font.bold = True

                doc.add_page_break()

                # ════════ Datos calculados para el reporte (fuerzas del sidebar) ════════
                cc_rows = []
                for (caso, joint), r in resultados.items():
                    if r['frf_op'] <= 0: continue
                    k_est = F_REF_N / r['frf_op']
                    cc_rows.append({'Caso': caso, 'Joint': joint, 'Dir': r['dir'],
                        'FRF_op (mm/T)': round(r['frf_op'], 4), 'K_est (N/mm)': round(k_est, 0),
                        'K_est/K_din': round(k_est/K_din, 2),
                        'Cumple 10:1': 'Sí' if k_est/K_din >= 10 else 'No'})
                df_cc_r = pd.DataFrame(cc_rows)

                cl_rows = []
                for (caso, joint), r in resultados.items():
                    F_op_cal = F_op_V if r['dir'] == 'Z' else F_op_H
                    F_rd_cal = F_rd_V if r['dir'] == 'Z' else F_rd_H
                    u_op = r['frf_op'] * F_op_cal / F_REF_N
                    vr_op = v_rms(u_op, f_op)
                    u_rd, f_rd_eff, _, _ = transitorio_uf(r, F_rd_cal, modo_fuerza_tr, U_gmm, F_REF_N, f_tr_lo, f_tr_hi, n_apoyos)
                    vr_rd = v_rms(u_rd, f_rd_eff)
                    cl_rows.append({'Caso': caso, 'Joint': joint, 'Dir': r['dir'],
                        'A_op (mm)': round(u_op, 5), 'vRMS_op (mm/s)': round(vr_op, 3),
                        'ISO_op': classify(vr_op, ISO_ZONES)[0],
                        'f_rd (Hz)': round(f_rd_eff, 2), 'A_rd (mm)': round(u_rd, 4),
                        'vRMS_rd (mm/s)': round(vr_rd, 2), 'ISO_rd': classify(vr_rd, ISO_ZONES)[0]})
                df_cl_r = pd.DataFrame(cl_rows)

                ratio_min_r = df_cc_r['K_est/K_din'].min() if not df_cc_r.empty else float('nan')
                dir_dom = (df_cc_r.loc[df_cc_r['FRF_op (mm/T)'].idxmax(), 'Dir']
                           if not df_cc_r.empty else '-')
                vop_max = df_cl_r['vRMS_op (mm/s)'].max() if not df_cl_r.empty else 0.0
                vrd_max = df_cl_r['vRMS_rd (mm/s)'].max() if not df_cl_r.empty else 0.0
                iso_op_w = classify(vop_max, ISO_ZONES)[0]
                iso_rd_w = classify(vrd_max, ISO_ZONES)[0]
                cc_cumple = (ratio_min_r >= 10) if ratio_min_r == ratio_min_r else True

                # ════════ 4. RESUMEN EJECUTIVO ════════
                add_h1("4. Resumen Ejecutivo")
                add_para(
                    f"Se realizaron análisis Steady-State en SAP2000 sobre los puntos de apoyo de la "
                    f"estructura de soporte, evaluando la condición de operación (régimen permanente a "
                    f"{f_op:.2f} Hz / {f_op_rpm:.0f} RPM) y la condición transitoria de partida/parada "
                    f"(f_rd ≈ {f_rd:.2f} Hz). Los principales hallazgos son:")
                add_para(
                    f"• Cortocircuito vibratorio: K_est/K_din = {ratio_min_r:.2f} "
                    f"{'≥' if cc_cumple else '<'} 10 (ACI 351.3R-04). "
                    f"{'Cumple.' if cc_cumple else 'No cumple: el aislador no logra desacoplar la máquina de la estructura.'}")
                add_para(f"• Dirección dominante: {dir_dom} (mayor FRF en operación).")
                add_para(
                    f"• Clasificación ISO 20816-3: en operación el caso más desfavorable alcanza "
                    f"v_RMS = {vop_max:.2f} mm/s ({iso_op_w}); en transitorio, "
                    f"v_RMS = {vrd_max:.2f} mm/s ({iso_rd_w}).")

                # ════════ 5. DEFINICIONES ════════
                add_h1("5. Definiciones")
                add_h2("Función de respuesta en frecuencia (FRF)")
                add_para("La FRF (Steady-State de SAP2000) es la relación entre el desplazamiento "
                         "complejo U(f) en un punto y la fuerza armónica F(f) aplicada, evaluada en cada "
                         "frecuencia del barrido.")
                add_h2("Ángulo de fase (φ)")
                add_para("Para un sistema de 1 GDL, φ = −90° exactamente en resonancia (r = 1). En "
                         "sistemas MDOF, la FRF superpone las contribuciones modales y φ(f) se obtiene de "
                         "las partes real e imaginaria de H(f).")
                add_h2("Cargas del fabricante y parámetros")
                df_par = pd.DataFrame([
                    {'Parámetro': 'Frecuencia de operación', 'Valor': f"{f_op:.2f} Hz ({f_op_rpm:.0f} RPM)"},
                    {'Parámetro': 'Frecuencia transitorio f_rd', 'Valor': f"{f_rd:.2f} Hz"},
                    {'Parámetro': 'Zona de exclusión', 'Valor': f"{f_excl_lo:.2f} – {f_excl_hi:.2f} Hz"},
                    {'Parámetro': 'K_din aislador', 'Valor': f"{K_din:.0f} N/mm"},
                    {'Parámetro': 'F_op (V / H) por apoyo', 'Valor': f"{F_op_V} / {F_op_H} N"},
                    {'Parámetro': 'F_rd (V / H) por apoyo', 'Valor': f"{F_rd_V} / {F_rd_H} N"},
                    {'Parámetro': 'Fuerza de referencia', 'Valor': f"{F_REF_N:.0f} N"},
                ])
                add_df_table(df_par, fs=10)

                # ════════ 6. ANÁLISIS DE VIBRACIONES ════════
                add_h1("6. Análisis de Vibraciones")
                add_h2("6.1 Cortocircuito Vibratorio y Rigidez Relativa")
                add_para("Criterio 10:1 (ACI 351.3R-04, §3.3): la rigidez dinámica de la estructura "
                         "(K_est = F_ref / FRF_op) debe ser al menos 10 veces la del aislador (K_din) para "
                         "que el aislamiento sea efectivo. El caso gobernante es el de mayor FRF_op.")
                add_para(f"Resultado: K_est/K_din mínimo = {ratio_min_r:.2f}. "
                         f"{'Cumple.' if cc_cumple else 'No cumple — cortocircuito vibratorio confirmado.'}")
                add_df_table(df_cc_r)

                add_h2("6.2 Análisis de Resonancia: Zona de Exclusión y Fase")
                add_para(f"Zona de exclusión ACI 351.3R-04: {f_excl_lo:.2f} ≤ fn ≤ {f_excl_hi:.2f} Hz "
                         f"(0.8·fop – 1.2·fop). La resonancia se confirma cuando el peak de |H(f)| cae en "
                         f"la zona y la fase está próxima a ±90°.")
                add_df_table(df_res)

                add_h2("6.3 Distorsión Angular Dinámica")
                dist_rows = []
                for caso in sorted(set(r['caso'] for r in resultados.values())):
                    fases = [(r['joint'], r['fase_op']) for (c, j), r in resultados.items() if c == caso]
                    maxd = 0.0
                    for i in range(len(fases)):
                        for jx in range(i+1, len(fases)):
                            d = abs(fases[i][1] - fases[jx][1]); d = 360-d if d > 180 else d
                            maxd = max(maxd, d)
                    dist_rows.append({'Caso': caso, 'N° apoyos': len(fases),
                                      'Δφ máx entre apoyos (°)': round(maxd, 1)})
                add_para("El desfase entre apoyos de una misma máquina impone torsión dinámica cíclica "
                         "sobre el chasis. Se reporta el diferencial de fase máximo entre apoyos por caso "
                         "(en operación):")
                add_df_table(pd.DataFrame(dist_rows), fs=10)

                add_h2("6.4 Criterios de Amplitud Admisible")
                add_para("Las amplitudes reales se obtienen como A = FRF · F_real / F_ref, y la velocidad "
                         "eficaz v_RMS = (A·2πf)/√2 se clasifica según ISO 20816-3. La operación se evalúa "
                         "en fop (con el peak estructural buscado dentro de la zona de exclusión); el "
                         "transitorio se evalúa como el peor caso a lo largo de la ventana de barrido "
                         f"({min(f_tr_lo, f_tr_hi):.2f}–{max(f_tr_lo, f_tr_hi):.2f} Hz) que el equipo "
                         "recorre en partida/parada.")
                if modo_fuerza_tr.startswith("Curva"):
                    add_para("Nota — modelo de fuerza transitoria por desbalance: la fuerza varía con la "
                             "frecuencia según F(f) = m·e·ω² ∝ f² (desbalance rotativo, válido para "
                             "equipos centrífugos). En consecuencia, la fuerza no es constante en el "
                             "barrido: es mínima en el cruce del aislador (f_rd, baja frecuencia) y "
                             "máxima cerca de operación. El peor caso del transitorio tiende, por tanto, "
                             "a ubicarse en la zona alta del barrido y no en f_rd.")
                else:
                    add_para("Nota — modelo de fuerza transitoria de valor fijo: se aplica la fuerza "
                             "dinámica máxima del fabricante (F_rd) como valor constante a lo largo de "
                             "todo el barrido (criterio conservador para el cruce de resonancia).")
                add_df_table(df_cl_r)

                # Figuras (matplotlib)
                if incluir_figs and resultados:
                    add_h2("Figura 6.1. FRF |U| por caso y nodo")
                    casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
                    joints_uniq = sorted(set(r['joint'] for r in resultados.values()))
                    n_c = min(len(casos_uniq), 3)
                    n_r = math.ceil(len(casos_uniq) / n_c)

                    fig_mpl, axes = plt.subplots(n_r, n_c, figsize=(14, 4*n_r))
                    if n_r*n_c == 1: axes = np.array([[axes]])
                    elif n_r == 1 or n_c == 1: axes = axes.reshape(n_r, n_c)

                    for ci, caso in enumerate(casos_uniq):
                        ax = axes[ci//n_c][ci%n_c]
                        grupo = {k: v for k, v in resultados.items() if k[0] == caso}
                        for (c, joint), r in grupo.items():
                            col_f = color_joint(joints_uniq, joint)
                            ax.plot(r['freqs'], r['frf_mm'],
                                    color=col_f[:3], lw=1.8, label=joint)
                            ax.plot(r['f_pk'], r['frf_pk'], 'o', color=col_f[:3], ms=6)
                        ax.axvspan(f_excl_lo, f_excl_hi, color='orange', alpha=0.15)
                        ax.axvline(f_op, color='red', ls='--', lw=1.5, alpha=0.7)
                        ax.set_title(caso, fontsize=10, fontweight='bold')
                        ax.set_xlabel('Frecuencia (Hz)', fontsize=9)
                        ax.set_ylabel('FRF (mm/Ton)', fontsize=9)
                        ax.legend(fontsize=7.5, ncol=2)
                        ax.grid(True, alpha=0.25, ls='--')
                        ax.set_xlim(df_raw['Freq'].min()*0.95, df_raw['Freq'].max()*1.05)
                        ax.set_ylim(bottom=0)

                    plt.tight_layout()
                    buf = io.BytesIO()
                    fig_mpl.savefig(buf, format='png', dpi=160, bbox_inches='tight')
                    buf.seek(0)
                    plt.close(fig_mpl)

                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_img.add_run()
                    run.add_picture(buf, width=Cm(15))

                    # --- Figuras de verificación de amplitudes (Richart / Blake / ISO) ---
                    def add_fig_mpl(fig, titulo, caption=None):
                        add_h2(titulo)
                        if caption:
                            add_para(caption)
                        b = io.BytesIO()
                        fig.savefig(b, format='png', dpi=160, bbox_inches='tight')
                        b.seek(0); plt.close(fig)
                        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pi.add_run().add_picture(b, width=Cm(15))

                    # Puntos (Op / Tr) por nodo, coloreados como en la app
                    joints_all = sorted(set(j for (_, j) in resultados.keys()))
                    def _hex_joint(j):
                        return '#%02x%02x%02x' % tuple(
                            int(x*255) for x in color_joint(joints_all, j)[:3])
                    puntos_cl = []
                    for _, row in df_cl_r.iterrows():
                        col_j = _hex_joint(row['Joint'])
                        if modo_cond in ["Operación", "Ambas"]:
                            puntos_cl.append((f_op*60, row['A_op (mm)'], str(row['Joint']), col_j, 'o'))
                        if modo_cond in ["Transitorio", "Ambas"]:
                            puntos_cl.append((row['f_rd (Hz)']*60, row['A_rd (mm)'], str(row['Joint']), col_j, '^'))

                    cap_pts = ("Marcadores: ○ = operación (f_op), △ = transitorio (peor caso del "
                               "tramo). Banda naranja = zona de exclusión; línea roja = f_op.")
                    add_fig_mpl(
                        fig_clasif_loglog_mpl(RICHART_LINES, puntos_cl, f_excl_lo, f_excl_hi,
                                              f_op, "Richart Fig. 10-1 — Amplitud vs Frecuencia",
                                              ylim=(10**-4.2, 10**1.2)),
                        "Figura 6.2. Verificación de amplitudes — Richart", cap_pts)
                    add_fig_mpl(
                        fig_clasif_loglog_mpl(BLAKE_LINES, puntos_cl, f_excl_lo, f_excl_hi,
                                              f_op, "Blake Fig. 10-2 — Amplitud vs Frecuencia",
                                              ylim=(10**-4.5, 10**0.6)),
                        "Figura 6.3. Verificación de amplitudes — Blake", cap_pts)
                    add_fig_mpl(
                        fig_iso_barras_mpl(df_cl_r, iso_a, iso_b, iso_c,
                                           "ISO 20816-3 — v_RMS en operación por nodo"),
                        "Figura 6.4. Verificación de amplitudes — ISO 20816-3")

                # ════════ 7. MEDIDAS DE CONTROL ════════
                add_h1("7. Medidas de Control")
                add_para("Según el resultado del criterio 10:1 y la clasificación de amplitudes, se "
                         "plantean dos estrategias de mitigación:")
                add_h2("Opción A — Refuerzo Estructural (High-Tuned)")
                add_para("Aumentar la rigidez dinámica de la estructura de soporte (K_est) hasta cumplir "
                         "K_est ≥ 10·K_din, llevando las frecuencias naturales locales por sobre la zona "
                         "de exclusión. Implica rigidizar vigas/atiesadores de los apoyos.")
                add_h2("Opción B — Cambio de Aislamiento (Low-Tuned)")
                add_para("Adoptar aisladores más blandos (menor K_din) para bajar la frecuencia natural "
                         "del sistema masa-aislador muy por debajo de la operación, mejorando la "
                         "transmisibilidad. Requiere verificar deflexión estática y estabilidad.")

                # Guardar y descargar
                buf_doc = io.BytesIO()
                doc.save(buf_doc)
                buf_doc.seek(0)

                st.success("✅ Reporte generado exitosamente")
                st.download_button(
                    label="⬇️ Descargar Reporte (.docx)",
                    data=buf_doc,
                    file_name=f"Auditoria_Dinamica_{revision}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ── Footer ────────────────────────────────────────────────────────────────────
st.divider()
st.caption(
    "Analizador Dinámico — ACI 351.3R-04 · ISO 20816-3:2022 · Richart (1962) · Blake (1964) "
    "| Datos: SAP2000 v27.1.0 (Kgf, cm, C) | F_ref = 1 Ton = 9810 N"
)
