"""
Analizador Dinámico de Plataformas de Soporte para Equipos Rotativos
======================================================================
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io, math, warnings
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as q
warnings.filterwarnings("ignore")

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURACIÓN
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Analizador Dinámico — Plataformas de Soporte",
    page_icon="📊", layout="wide", initial_sidebar_state="expanded"
)

st.markdown("""
<style>
[data-testid="stSidebar"] { background: #1B2A4A; }
[data-testid="stSidebar"] * { color: #E2E8F0 !important; }
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stNumberInput label { color: #94A3B8 !important; font-size:13px; }

/* Corrección para que se vea el texto en los inputs numéricos */
[data-testid="stSidebar"] input,
[data-baseweb="input"] input,
div[data-baseweb="select"] * { color: #1B2A4A !important; -webkit-text-fill-color: #1B2A4A !important; }
[data-baseweb="input"] > div { background-color: #FFFFFF !important; border: 1px solid #94A3B8 !important; border-radius: 4px; }

h1 { color: #1B2A4A; font-family: Arial; }
h2 { color: #2D5F8A; font-family: Arial; border-bottom: 2px solid #2D5F8A; padding-bottom:4px; }
h3 { color: #334155; font-family: Arial; }
.metric-card { background:#EFF6FF; border-left:4px solid #2563EB;
               padding:10px 14px; border-radius:4px; margin:6px 0; }
.warn-box    { background:#FFFBEB; border-left:4px solid #D97706;
               padding:10px 14px; border-radius:4px; }
.crit-box    { background:#FEF2F2; border-left:4px solid #DC2626;
               padding:10px 14px; border-radius:4px; }
.ok-box      { background:#F0FDF4; border-left:4px solid #16A34A;
               padding:10px 14px; border-radius:4px; }
.stDataFrame { font-size: 12px; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTES Y ZONAS DE CLASIFICACIÓN
# ══════════════════════════════════════════════════════════════════════════════
F_REF_N = 9810.0  # N/Ton (fuerza de referencia SAP2000)

RICHART_ZONES = [
    (0.787,  "No perceptible",          "#1a9850"),
    (1.854,  "Apenas perceptible",      "#66bd63"),
    (5.309,  "Fácilmente perceptible",  "#fee08b"),
    (10.640, "Molesta",                 "#fdae61"),
    (21.290, "Severa (personas)",       "#f46d43"),
    (53.090, "Límite máquinas",         "#d73027"),
    (127.00, "Precaución estructuras",  "#9e0142"),
    (1e9,    "Peligro estructuras",     "#67001f"),
]
BLAKE_ZONES = [
    (0.360,  "A — Sin fallas",         "#1a9850"),
    (3.590,  "B — Fallas menores",     "#66bd63"),
    (35.90,  "C — Defectuoso",         "#fee08b"),
    (107.7,  "D — Falla inminente",    "#f46d43"),
    (1e9,    "E — Peligroso",          "#d73027"),
]
ISO_ZONES = [
    (1.12, "Zona A — Aceptable",        "#1a9850"),
    (2.80, "Zona B — Normal",           "#66bd63"),
    (7.10, "Zona C — Alarma",           "#fee08b"),
    (1e9,  "Zona D — Detener equipo",   "#f46d43"),
]

def classify(v, zones):
    for v_max, label, color in zones:
        if v <= v_max:
            return label, color
    return zones[-1][1], zones[-1][2]

def v_peak(amp_mm, f_hz):
    return amp_mm * 2 * math.pi * f_hz

def v_rms(amp_mm, f_hz):
    return v_peak(amp_mm, f_hz) / math.sqrt(2)

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR — PARÁMETROS
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Parámetros del Proyecto")
    st.divider()

    st.markdown("**Equipo**")
    f_op = st.number_input("Frecuencia de operación (Hz)", value=48.33, step=0.01, format="%.2f")
    st.caption(f"= {f_op*60:.0f} cpm")

    st.divider()
    st.markdown("**Cargas del fabricante (N/pata = total ÷ 4)**")
    col1, col2 = st.columns(2)
    with col1:
        F_op_V  = st.number_input("F_op vertical",     value=425,  step=50)
        F_op_H  = st.number_input("F_op horizontal",   value=425,  step=50)
    with col2:
        F_rd_V  = st.number_input("F_rd vertical",     value=5000, step=100)
        F_rd_H  = st.number_input("F_rd horizontal",   value=2500, step=100)

    st.divider()
    st.markdown("**Aislador**")
    K_din = st.number_input("K_din aislador (N/mm)", value=5600.0, step=100.0, format="%.0f")
    st.caption("Novibra RAEM2500: 5600 N/mm")

    st.divider()
    st.markdown("**Zona de exclusión**")
    z_lo = st.number_input("Factor inferior", value=0.8, step=0.05, format="%.2f")
    z_hi = st.number_input("Factor superior", value=1.2, step=0.05, format="%.2f")
    f_excl_lo = z_lo * f_op
    f_excl_hi = z_hi * f_op
    st.caption(f"= {f_excl_lo:.2f} – {f_excl_hi:.2f} Hz")

    st.divider()
    modo_cond = st.selectbox("Condición a graficar",
                             ["Operación", "Run-Down", "Ambas"], index=0)

# ══════════════════════════════════════════════════════════════════════════════
# HEADER PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
st.title("📊 Analizador Dinámico — Plataformas de Soporte para Equipos Rotativos")
st.markdown(
    f"**Frecuencia de operación:** {f_op} Hz ({f_op*60:.0f} cpm) | "
    f"**Zona de exclusión ACI 351.3R:** {f_excl_lo:.2f} – {f_excl_hi:.2f} Hz | "
    f"**K_din aislador:** {K_din:.0f} N/mm"
)
st.divider()

# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════
tab_data, tab_frf, tab_fase, tab_class, tab_dist, tab_report = st.tabs([
    "📥 Datos SAP2000",
    "📈 FRF — Amplitud",
    "🔄 Análisis de Fase",
    "🗂️ Clasificación Richart / Blake / ISO",
    "🔀 Distorsión Angular",
    "📄 Reporte"
])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1: DATOS DE ENTRADA
# ══════════════════════════════════════════════════════════════════════════════
with tab_data:
    st.header("Datos de entrada — SAP2000 Steady-State")

    col_up, col_ex = st.columns([2, 1])
    with col_up:
        uploaded = st.file_uploader(
            "Cargar archivo Excel exportado de SAP2000 (.xlsx)",
            type=["xlsx", "csv"],
            help="Exportar desde SAP2000: Display → Show Tables → Analysis Results → Joint Displacements"
        )
    with col_ex:
        st.markdown("**Formato esperado de columnas:**")
        st.code("Joint | OutputCase | CaseType | StepType | Freq | U1 | U2 | U3", language="")
        st.caption("StepTypes: Mag at Freq / Real at Freq / Imag at Freq")

    with st.expander("O ingresar datos directamente (Formato CSV):"):
        sample_csv = """Joint,OutputCase,CaseType,StepType,Freq,U1,U2,U3,R1,R2,R3
C11,STST_X,LinSteadyState,Mag at Freq,48.33,2.514E-02,1.2E-03,8.5E-04,0,0,0
C11,STST_X,LinSteadyState,Real at Freq,48.33,2.480E-02,1.1E-03,8.0E-04,0,0,0
C11,STST_X,LinSteadyState,Imag at Freq,48.33,-4.2E-03,-2.1E-04,-2.0E-04,0,0,0
C14,STST_X,LinSteadyState,Mag at Freq,57.30,9.787E-02,3.1E-03,9.2E-04,0,0,0
C14,STST_X,LinSteadyState,Real at Freq,57.30,5.9E-03,2.8E-04,8.5E-05,0,0,0
C14,STST_X,LinSteadyState,Imag at Freq,57.30,-9.784E-02,-3.1E-03,-9.2E-04,0,0,0"""

        manual_data = st.text_area(
            "Pegar datos CSV (Joint, OutputCase, CaseType, StepType, Freq, U1, U2, U3, ...)",
            value=sample_csv, height=200
        )

    @st.cache_data
    def load_data(file_bytes, file_name, manual_text):
        try:
            if file_bytes:
                if file_name.endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(file_bytes))
                else:
                    df = pd.read_excel(io.BytesIO(file_bytes), skiprows=2)
                    df.columns = ['Joint','OutputCase','CaseType','StepType',
                                  'Freq','U1','U2','U3','R1','R2','R3']
            else:
                df = pd.read_csv(io.StringIO(manual_text))
            for col in ['Freq','U1','U2','U3']:
                df[col] = pd.to_numeric(df[col], errors='coerce')
            return df.dropna(subset=['Freq'])
        except Exception as e:
            st.error(f"Error al cargar datos: {e}")
            return pd.DataFrame()

    file_bytes = uploaded.read() if uploaded else None
    file_name  = uploaded.name  if uploaded else ""
    df_raw = load_data(file_bytes, file_name, manual_data)

    if not df_raw.empty:
        st.success(f"✅ {len(df_raw)} filas cargadas | "
                   f"Juntas: {sorted(df_raw['Joint'].unique())} | "
                   f"Casos: {list(df_raw['OutputCase'].unique())} | "
                   f"StepTypes: {list(df_raw['StepType'].unique())}")

        col_fil1, col_fil2 = st.columns(2)
        with col_fil1:
            casos_disp = sorted(df_raw['OutputCase'].unique())
            casos_sel  = st.multiselect("Casos a analizar", casos_disp, default=casos_disp)
        with col_fil2:
            juntas_disp = sorted(df_raw['Joint'].unique())
            juntas_sel  = st.multiselect("Juntas a analizar", juntas_disp, default=juntas_disp)

        df_raw = df_raw[df_raw['OutputCase'].isin(casos_sel) &
                        df_raw['Joint'].isin(juntas_sel)]
        st.dataframe(df_raw.head(30), use_container_width=True, height=250)
    else:
        st.warning("Cargue datos para continuar.")
        casos_sel = []; juntas_sel = []; df_raw = pd.DataFrame()

# ══════════════════════════════════════════════════════════════════════════════
# PROCESAMIENTO CENTRAL
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def procesar(df_raw_json, f_op_, f_excl_lo_, f_excl_hi_):
    if not df_raw_json:
        return {}, pd.DataFrame()
    df = pd.read_json(io.StringIO(df_raw_json))
    mag_df  = df[df['StepType']=='Mag at Freq'].copy()
    real_df = df[df['StepType']=='Real at Freq'].copy()
    imag_df = df[df['StepType']=='Imag at Freq'].copy()

    # Mapeo case → (col_resp, dir_label)
    case_dir = {}
    for caso in df['OutputCase'].unique():
        caso_u = caso.upper()
        if '_X' in caso_u or caso_u.endswith('X'):
            case_dir[caso] = ('U1', 'X')
        elif '_Y' in caso_u or caso_u.endswith('Y'):
            case_dir[caso] = ('U2', 'Y')
        elif '_Z' in caso_u or caso_u.endswith('Z'):
            case_dir[caso] = ('U3', 'Z')
        else:
            case_dir[caso] = ('U1', caso)

    resultados = {}  # {(caso, junta): {...}}

    for caso in df['OutputCase'].unique():
        col_resp, dir_lbl = case_dir.get(caso, ('U1', caso))
        for junta in df['Joint'].unique():
            mj = mag_df[(mag_df['OutputCase']==caso)&(mag_df['Joint']==junta)].sort_values('Freq')
            rj = real_df[(real_df['OutputCase']==caso)&(real_df['Joint']==junta)].sort_values('Freq')
            ij = imag_df[(imag_df['OutputCase']==caso)&(imag_df['Joint']==junta)].sort_values('Freq')
            if mj.empty: continue

            freqs   = mj['Freq'].values
            frf_mm  = np.abs(mj[col_resp].values) * 10  # cm → mm, /Ton
            fase_deg= np.degrees(np.arctan2(
                ij[col_resp].values if not ij.empty else np.zeros_like(freqs),
                rj[col_resp].values if not rj.empty else np.zeros_like(freqs)
            ))

            mask = (freqs >= 40) & (freqs <= 80)
            if not mask.any(): mask = np.ones(len(freqs), bool)

            idx_pk  = np.argmax(frf_mm[mask])
            f_pk    = freqs[mask][idx_pk]
            frf_pk  = frf_mm[mask][idx_pk]
            fase_pk = fase_deg[mask][idx_pk]

            idx_op  = np.argmin(np.abs(freqs - f_op_))
            frf_op  = frf_mm[idx_op]
            fase_op = fase_deg[idx_op]

            en_zona = f_excl_lo_ <= f_pk <= f_excl_hi_
            dist_90 = abs(abs(fase_pk) - 90)

            if dist_90 < 15 and en_zona and frf_pk > 0.05:
                diagnostico = "Resonancia confirmada"
            elif dist_90 < 30 and en_zona:
                diagnostico = "Probable"
            elif en_zona:
                diagnostico = "Posible"
            else:
                diagnostico = "Sin resonancia en zona"

            resultados[(caso, junta)] = {
                'caso': caso, 'junta': junta, 'dir': dir_lbl,
                'freqs': freqs.tolist(), 'frf_mm': frf_mm.tolist(),
                'fase_deg': fase_deg.tolist(),
                'f_pk': f_pk, 'frf_pk': frf_pk, 'fase_pk': fase_pk,
                'frf_op': frf_op, 'fase_op': fase_op,
                'en_zona': en_zona, 'dist_90': dist_90,
                'diagnostico': diagnostico,
            }

    # DataFrame resumen
    rows = []
    for (caso, junta), r in resultados.items():
        rows.append({
            'Caso': caso, 'Junta': junta, 'Dir': r['dir'],
            'f_peak (Hz)': round(r['f_pk'], 2),
            'FRF_peak (mm/T)': round(r['frf_pk'], 4),
            'FRF_op (mm/T)': round(r['frf_op'], 4),
            'Fase_peak (°)': round(r['fase_pk'], 1),
            'En zona excl.': '✓ Sí' if r['en_zona'] else 'No',
            'Diagnóstico': r['diagnostico'],
        })
    return resultados, pd.DataFrame(rows)

if not df_raw.empty:
    resultados, df_res = procesar(
        df_raw.to_json(), f_op, f_excl_lo, f_excl_hi
    )
else:
    resultados, df_res = {}, pd.DataFrame()

# ── Colores por junta ─────────────────────────────────────────────────────────
PALETA = plt.cm.tab20.colors
def color_junta(juntas_list, junta):
    idx = sorted(set(juntas_list)).index(junta) % len(PALETA)
    return PALETA[idx]

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2: FRF — AMPLITUD
# ══════════════════════════════════════════════════════════════════════════════
with tab_frf:
    st.header("FRF — Amplitud |U| por caso y junta")
    if not resultados:
        st.info("Cargue datos en la pestaña 'Datos SAP2000' para ver los gráficos.")
    else:
        casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
        juntas_uniq = sorted(set(r['junta'] for r in resultados.values()))

        n_cols = st.slider("Columnas por fila", 1, 3, min(len(casos_uniq), 3))
        fig_frf = make_subplots(
            rows=math.ceil(len(casos_uniq)/n_cols), cols=n_cols,
            subplot_titles=casos_uniq, shared_yaxes=False
        )

        for ci, caso in enumerate(casos_uniq):
            row_p = ci // n_cols + 1; col_p = ci % n_cols + 1
            grupo = {k: v for k, v in resultados.items() if k[0] == caso}
            for (c, junta), r in grupo.items():
                col_hex = '#%02x%02x%02x' % tuple(
                    int(x*255) for x in color_junta(juntas_uniq, junta)[:3])
                fig_frf.add_trace(go.Scatter(
                    x=r['freqs'], y=r['frf_mm'],
                    name=junta, legendgroup=junta,
                    showlegend=(ci == 0),
                    line=dict(color=col_hex, width=2),
                    mode='lines',
                    hovertemplate=f"<b>{junta}</b><br>f=%{{x:.2f}} Hz<br>FRF=%{{y:.4f}} mm/T<extra></extra>"
                ), row=row_p, col=col_p)
                # Marcar peak
                fig_frf.add_trace(go.Scatter(
                    x=[r['f_pk']], y=[r['frf_pk']],
                    mode='markers', showlegend=False,
                    marker=dict(color=col_hex, size=8, symbol='circle'),
                    hovertemplate=f"<b>{junta} peak</b><br>f={r['f_pk']:.1f} Hz<br>{r['frf_pk']:.4f} mm/T<extra></extra>"
                ), row=row_p, col=col_p)

            # Zona de exclusión
            fig_frf.add_vrect(
                x0=f_excl_lo, x1=f_excl_hi,
                fillcolor="orange", opacity=0.12, line_width=0,
                row=row_p, col=col_p
            )
            fig_frf.add_vline(
                x=f_op, line_dash="dash", line_color="red", opacity=0.7,
                row=row_p, col=col_p
            )

        fig_frf.update_layout(
            height=380*math.ceil(len(casos_uniq)/n_cols),
            title_text="FRF |U| — Desplazamiento por unidad de fuerza (mm/Ton)",
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            font=dict(family="Arial"),
            hovermode="x unified"
        )
        fig_frf.update_xaxes(title_text="Frecuencia (Hz)", range=[
            df_raw['Freq'].min()*0.95 if not df_raw.empty else 30,
            df_raw['Freq'].max()*1.05 if not df_raw.empty else 80
        ])
        fig_frf.update_yaxes(title_text="|U| (mm/Ton)", rangemode="tozero")
        st.plotly_chart(fig_frf, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3: ANÁLISIS DE FASE
# ══════════════════════════════════════════════════════════════════════════════
with tab_fase:
    st.header("Ángulo de fase φ — Criterio de resonancia (|φ| ≈ 90°)")
    if not resultados:
        st.info("Cargue datos para ver el análisis de fase.")
    else:
        casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
        juntas_uniq = sorted(set(r['junta'] for r in resultados.values()))
        n_cols = min(len(casos_uniq), 3)

        fig_fase = make_subplots(
            rows=math.ceil(len(casos_uniq)/n_cols), cols=n_cols,
            subplot_titles=casos_uniq
        )
        for ci, caso in enumerate(casos_uniq):
            row_p = ci // n_cols + 1; col_p = ci % n_cols + 1
            grupo = {k: v for k, v in resultados.items() if k[0] == caso}
            for (c, junta), r in grupo.items():
                col_hex = '#%02x%02x%02x' % tuple(
                    int(x*255) for x in color_junta(juntas_uniq, junta)[:3])
                fig_fase.add_trace(go.Scatter(
                    x=r['freqs'], y=r['fase_deg'],
                    name=junta, legendgroup=junta,
                    showlegend=(ci == 0),
                    line=dict(color=col_hex, width=2),
                    mode='lines',
                    hovertemplate=f"<b>{junta}</b><br>f=%{{x:.2f}} Hz<br>φ=%{{y:.1f}}°<extra></extra>"
                ), row=row_p, col=col_p)
                fig_fase.add_trace(go.Scatter(
                    x=[r['f_pk']], y=[r['fase_pk']],
                    mode='markers+text',
                    text=[f"{r['fase_pk']:.0f}°"],
                    textposition="top right",
                    showlegend=False,
                    marker=dict(color=col_hex, size=8),
                    textfont=dict(size=9, color=col_hex)
                ), row=row_p, col=col_p)

            fig_fase.add_vrect(
                x0=f_excl_lo, x1=f_excl_hi,
                fillcolor="orange", opacity=0.12, line_width=0,
                row=row_p, col=col_p
            )
            fig_fase.add_vline(x=f_op, line_dash="dash", line_color="red", opacity=0.7, row=row_p, col=col_p)
            for y_ref in [-90, 90]:
                fig_fase.add_hline(y=y_ref, line_dash="dot", line_color="purple",
                                   opacity=0.6, row=row_p, col=col_p)

        fig_fase.update_layout(
            height=380*math.ceil(len(casos_uniq)/n_cols),
            title_text="Ángulo de fase φ — Líneas ±90° = criterio de resonancia",
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            font=dict(family="Arial"),
        )
        fig_fase.update_yaxes(range=[-220, 220], title_text="φ (°)")
        fig_fase.update_xaxes(title_text="Frecuencia (Hz)")
        st.plotly_chart(fig_fase, use_container_width=True)

        # Tabla diagnóstico
        st.subheader("Resumen de Diagnóstico de Resonancia")
        if not df_res.empty:
            def color_diag(val):
                if "confirmada" in str(val): return "background-color:#FEE2E2; color:#C0392B; font-weight:bold"
                if "Probable"   in str(val): return "background-color:#FEF3C7; color:#D97706"
                if "Posible"    in str(val): return "background-color:#FFFBEB; color:#92400E"
                return "background-color:#F0FDF4; color:#16A34A"
            styled = df_res.style.map(color_diag, subset=['Diagnóstico'])
            st.dataframe(styled, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
# TAB 4: CLASIFICACIÓN RICHART / BLAKE / ISO
# ══════════════════════════════════════════════════════════════════════════════
with tab_class:
    st.header("Clasificación de amplitudes — Richart, Blake e ISO 20816-3")
    if not resultados:
        st.info("Cargue datos para ver la clasificación.")
    else:
        st.markdown("**Configurar fuerzas reales para calcular amplitudes:**")
        col_a, col_b, col_c, col_d = st.columns(4)
        with col_a: F_op_V_c  = st.number_input("F op V (N)", value=F_op_V,  key="fopv_c")
        with col_b: F_op_H_c  = st.number_input("F op H (N)", value=F_op_H,  key="foph_c")
        with col_c: F_rd_V_c  = st.number_input("F rd V (N)", value=F_rd_V,  key="frdv_c")
        with col_d: F_rd_H_c  = st.number_input("F rd H (N)", value=F_rd_H,  key="frdh_c")

        # Construir tabla de resultados de clasificación
        rows_class = []
        for (caso, junta), r in resultados.items():
            dir_lbl = r['dir']
            F_op = F_op_V_c if dir_lbl == 'Z' else F_op_H_c
            F_rd = F_rd_V_c if dir_lbl == 'Z' else F_rd_H_c

            # Operación: FRF a f_op
            u_op  = r['frf_op'] * F_op / F_REF_N
            vp_op = v_peak(u_op, f_op)
            vr_op = v_rms(u_op,  f_op)
            r_op_lbl, r_op_col = classify(vp_op, RICHART_ZONES)
            b_op_lbl, b_op_col = classify(vp_op, BLAKE_ZONES)
            i_op_lbl, i_op_col = classify(vr_op, ISO_ZONES)

            # Run-down: FRF en peak
            u_rd  = r['frf_pk'] * F_rd / F_REF_N
            vp_rd = v_peak(u_rd, r['f_pk'])
            vr_rd = v_rms(u_rd,  r['f_pk'])
            r_rd_lbl, _ = classify(vp_rd, RICHART_ZONES)
            b_rd_lbl, _ = classify(vp_rd, BLAKE_ZONES)
            i_rd_lbl, _ = classify(vr_rd, ISO_ZONES)

            rows_class.append({
                'Caso': caso, 'Junta': junta, 'Dir': dir_lbl,
                'f_peak (Hz)': round(r['f_pk'], 1),
                # Operación
                'A_op (mm)': round(u_op, 5),
                'v_op (mm/s)': round(vp_op, 3),
                'vRMS_op (mm/s)': round(vr_op, 3),
                'Richart_op': r_op_lbl,
                'Blake_op': b_op_lbl,
                'ISO_op': i_op_lbl,
                # Run-Down
                'A_rd (mm)': round(u_rd, 4),
                'v_rd (mm/s)': round(vp_rd, 2),
                'vRMS_rd (mm/s)': round(vr_rd, 2),
                'Richart_rd': r_rd_lbl,
                'ISO_rd': i_rd_lbl,
            })

        df_class = pd.DataFrame(rows_class)

        # Mostrar con colores
        col_crit = ['Richart_op','Blake_op','ISO_op','Richart_rd','ISO_rd']
        def color_class(val):
            val = str(val)
            if any(x in val for x in ["Peligro","Zona D","Severa","E — "]):
                return "background-color:#FEE2E2;color:#C0392B;font-weight:bold"
            if any(x in val for x in ["Precaución","Límite","Zona C","C — ","D — ","Molesta"]):
                return "background-color:#FEF3C7;color:#D97706"
            if any(x in val for x in ["Fácilmente","B — ","Zona B"]):
                return "background-color:#FFFBEB"
            if any(x in val for x in ["No perceptible","Zona A","A — ","Apenas"]):
                return "background-color:#F0FDF4;color:#166534"
            return ""
        
        st.dataframe(
            df_class.style.map(color_class, subset=col_crit),
            use_container_width=True, height=400
        )

        # Gráfico de barras v_RMS para ISO
        st.subheader("ISO 20816-3 — v_RMS por nodo y dirección")
        
        # Ajustar el tope Y dinámicamente según el valor máximo de los datos
        max_vrms = df_class['vRMS_op (mm/s)'].max() if not df_class.empty else 0
        if pd.isna(max_vrms): max_vrms = 0
        iso_top_y = max(max_vrms * 1.2, 10.0) # Asegura al menos ver hasta 10 mm/s
        
        fig_iso = go.Figure()
        ISO_LIMS  = [1.12, 2.80, 7.10]
        ISO_NAMES = ["Zona A (≤1.12)", "Zona B (≤2.80)", "Zona C (≤7.10)"]
        ISO_COLS  = ["#1a9850",        "#fee08b",         "#f46d43"]
        for y_val, name, col_ in zip(ISO_LIMS, ISO_NAMES, ISO_COLS):
            fig_iso.add_hline(y=y_val, line_dash="dash", line_color=col_,
                              annotation_text=name, annotation_position="right",
                              line_width=2)

        dir_colors = {"X":"#1f77b4","Y":"#e74c3c","Z":"#27ae60"}
        for dir_ in df_class['Dir'].unique():
            df_d = df_class[df_class['Dir']==dir_]
            labels = df_d['Junta'].astype(str) + "<br>(" + df_d['Caso'].astype(str) + ")"
            fig_iso.add_trace(go.Bar(
                name=f"Dir {dir_} — Operación",
                x=labels, y=df_d['vRMS_op (mm/s)'],
                marker_color=dir_colors.get(dir_, "#666"),
                opacity=0.85,
                hovertemplate="<b>%{x}</b><br>v_RMS = %{y:.3f} mm/s<extra></extra>"
            ))

        fig_iso.update_layout(
            title="ISO 20816-3 — Velocidad eficaz v_RMS en operación",
            yaxis_title="v_RMS (mm/s)", xaxis_title="Punto de apoyo",
            yaxis=dict(range=[0, iso_top_y]), # Aplicando el tope dinámico
            barmode="group", height=500,
            font=dict(family="Arial"),
            legend=dict(orientation="h", yanchor="bottom", y=-0.3)
        )
        fig_iso.add_hrect(y0=0, y1=1.12, fillcolor="#1a9850", opacity=0.06, line_width=0)
        fig_iso.add_hrect(y0=1.12, y1=2.80, fillcolor="#66bd63", opacity=0.06, line_width=0)
        fig_iso.add_hrect(y0=2.80, y1=7.10, fillcolor="#fee08b", opacity=0.08, line_width=0)
        fig_iso.add_hrect(y0=7.10, y1=iso_top_y, fillcolor="#f46d43", opacity=0.08, line_width=0)
        st.plotly_chart(fig_iso, use_container_width=True)

        # ══════════════════════════════════════════════════════════════════════
        # FORMATO EXACTO DE EJES LOGARÍTMICOS Y ESTILOS
        # ══════════════════════════════════════════════════════════════════════
        grid_style = dict(showgrid=True, gridwidth=1, gridcolor='rgba(255, 255, 255, 0.15)')
        minor_grid_style = dict(dtick="D1", showgrid=True, gridwidth=0.5, gridcolor='rgba(255, 255, 255, 0.05)')
        
        # Textos explícitos para no utilizar formato científico de Plotly
        X_TICK_VALS = [100, 1000, 10000]
        X_TICK_TEXT = ["100", "1000", "10000"]
        Y_TICK_VALS = [0.0001, 0.001, 0.01, 0.1, 1, 10, 100]
        Y_TICK_TEXT = ["0.0001", "0.001", "0.01", "0.1", "1", "10", "100"]

        # ----------------------------------------------------------------------
        # Gráfico Richart Fig 10-1 (Coordenadas Exactas extrapoladas a los bordes)
        # ----------------------------------------------------------------------
        st.subheader("Richart Fig. 10-1 — Puntos de análisis en diagrama log-log")
        fig_rich = go.Figure()

        # Coordenadas deducidas matemáticamente a partir de los puntos entregados 
        # y extrapoladas a los bordes (100 y 10000) manteniendo la pendiente física natural.
        richart_lines = {
            "Peligro estructuras":    {"x": [100, 580, 4100, 10000], "y": [10.08, 1.8796, 0.2921, 0.1250], "col": "#67001f", "dash": "solid"},
            "Precaución estructuras": {"x": [100, 640, 3500, 10000], "y": [9.113, 1.3462, 0.2311, 0.0785], "col": "#9e0142", "dash": "solid"},
            "Límite máquinas":        {"x": [100, 4600, 10000], "y": [2.54, 0.05461, 0.02506], "col": "#d73027", "dash": "solid"},
            "Severa (personas)":      {"x": [100, 540, 2600, 4200, 10000], "y": [14.46, 0.4572, 0.01829, 0.01067, 0.00400], "col": "#f46d43", "dash": "dash"},
            "Molesta (personas)":     {"x": [100, 270, 2250, 3800, 10000], "y": [2.846, 0.5588, 0.01778, 0.00965, 0.00312], "col": "#fdae61", "dash": "dash"},
            "Fácilmente perceptible": {"x": [100, 260, 5700, 10000], "y": [0.268, 0.10287, 0.00465, 0.00264], "col": "#fee08b", "dash": "dash"},
            "Apenas perceptible":     {"x": [100, 185, 3900, 10000], "y": [0.0846, 0.04699, 0.00254, 0.00103], "col": "#66bd63", "dash": "dash"},
            "No perceptible":         {"x": [100, 145, 1200, 10000], "y": [0.0336, 0.02286, 0.00254, 0.00028], "col": "#1a9850", "dash": "dash"}
        }

        for lbl, data in richart_lines.items():
            fig_rich.add_trace(go.Scatter(
                x=data["x"], y=data["y"], mode='lines', name=lbl,
                line=dict(color=data["col"], width=2, dash=data["dash"], shape='linear'),
                showlegend=True, hovertemplate=f"{lbl}<extra></extra>"
            ))

        plotted = set()
        for (caso, junta), r in resultados.items():
            dir_lbl = r['dir']
            F_op = F_op_V_c if dir_lbl == 'Z' else F_op_H_c
            F_rd = F_rd_V_c if dir_lbl == 'Z' else F_rd_H_c
            u_op = r['frf_op'] * F_op / F_REF_N
            u_rd = r['frf_pk'] * F_rd / F_REF_N
            col_ = '#%02x%02x%02x' % tuple(
                int(x*255) for x in color_junta(
                    sorted(set(rv['junta'] for rv in resultados.values())), junta)[:3])

            puntos_a_graficar = []
            if modo_cond in ["Operación", "Ambas"]:
                puntos_a_graficar.append((u_op, f_op, "Op", "circle"))
            if modo_cond in ["Run-Down", "Ambas"]:
                puntos_a_graficar.append((u_rd, r['f_pk'], "Rd", "triangle-up"))

            for u_val, f_val, cond_lbl, sym in puntos_a_graficar:
                key_ = f"{caso}_{junta}_{cond_lbl}"
                if key_ not in plotted and u_val > 0:
                    plotted.add(key_)
                    fig_rich.add_trace(go.Scatter(
                        x=[f_val*60], y=[u_val],
                        mode='markers',
                        name=f"{junta} {dir_lbl} {cond_lbl}",
                        marker=dict(symbol=sym, size=10, color=col_,
                                    line=dict(width=1.5, color="black")),
                        showlegend=False,
                        hovertemplate=(
                            f"<b>{junta} — Dir {dir_lbl} — {cond_lbl}</b><br>"
                            f"f = {f_val:.1f} Hz ({f_val*60:.0f} cpm)<br>"
                            f"A = {u_val:.5f} mm<extra></extra>")
                    ))

        fig_rich.add_vrect(x0=f_excl_lo*60, x1=f_excl_hi*60,
                           fillcolor="orange", opacity=0.15, line_width=0)
        fig_rich.add_vline(x=f_op*60, line_dash="dash", line_color="red", opacity=0.8)

        fig_rich.update_layout(
            title="Richart Fig. 10-1 — Amplitud vs Frecuencia",
            height=850, font=dict(family="Arial"), 
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            margin=dict(l=60, r=40, t=60, b=60), 
            plot_bgcolor='rgba(0,0,0,0)', 
            paper_bgcolor='rgba(0,0,0,0)'
        )
        
        fig_rich.update_xaxes(
            title="Frecuencia (cpm)", type="log", range=[2, 4], 
            tickmode="array", tickvals=X_TICK_VALS, ticktext=X_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        fig_rich.update_yaxes(
            title="Amplitud peak (mm)", type="log", range=[-4.2, 1.2], 
            tickmode="array", tickvals=Y_TICK_VALS, ticktext=Y_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        
        st.plotly_chart(fig_rich, use_container_width=True, config={'toImageButtonOptions': {'format': 'png', 'filename': 'Grafico_Richart', 'scale': 2}})

        # ----------------------------------------------------------------------
        # Gráfico Blake Fig 10-2 (Coordenadas Exactas de Quiebres Originales)
        # ----------------------------------------------------------------------
        st.subheader("Blake Fig. 10-2 — Criterios de severidad vibratoria")
        fig_blake = go.Figure()

        # Coordenadas deducidas matemáticamente a partir de los quiebres e interceptos suministrados,
        # extendiendo la pendiente principal ininterrumpida hasta los 100 cpm.
        blake_lines = {
            "E — Peligroso":       {"x": [100, 310, 1850, 3700, 10000], "y": [0.5842, 1.016, 0.2286, 0.12192, 0.01524], "col": "#d73027"},
            "D — Falla inminente": {"x": [100, 310, 1850, 3700, 10000], "y": [0.2286, 0.18288, 0.08636, 0.04191, 0.004826], "col": "#f46d43"},
            "C — Defectuoso":      {"x": [100, 310, 1850, 3700, 10000], "y": [0.09398, 0.0635, 0.02921, 0.013208, 0.00381], "col": "#fee08b"},
            "B — Fallas menores":  {"x": [100, 1850, 5900, 10000], "y": [0.0381, 0.009017, 0.00254, 0.001417], "col": "#66bd63"}
        }

        for lbl, data in blake_lines.items():
            fig_blake.add_trace(go.Scatter(
                x=data["x"], y=data["y"], mode='lines', name=lbl,
                line=dict(color=data["col"], width=2, dash='solid', shape='linear'),
                showlegend=True, hovertemplate=f"{lbl}<extra></extra>"
            ))

        plotted_blake = set()
        for (caso, junta), r in resultados.items():
            dir_lbl = r['dir']
            F_op = F_op_V_c if dir_lbl == 'Z' else F_op_H_c
            F_rd = F_rd_V_c if dir_lbl == 'Z' else F_rd_H_c
            u_op = r['frf_op'] * F_op / F_REF_N
            u_rd = r['frf_pk'] * F_rd / F_REF_N
            col_ = '#%02x%02x%02x' % tuple(
                int(x*255) for x in color_junta(
                    sorted(set(rv['junta'] for rv in resultados.values())), junta)[:3])

            puntos_a_graficar = []
            if modo_cond in ["Operación", "Ambas"]:
                puntos_a_graficar.append((u_op, f_op, "Op", "circle"))
            if modo_cond in ["Run-Down", "Ambas"]:
                puntos_a_graficar.append((u_rd, r['f_pk'], "Rd", "triangle-up"))

            for u_val, f_val, cond_lbl, sym in puntos_a_graficar:
                key_ = f"{caso}_{junta}_{cond_lbl}"
                if key_ not in plotted_blake and u_val > 0:
                    plotted_blake.add(key_)
                    fig_blake.add_trace(go.Scatter(
                        x=[f_val*60], y=[u_val],
                        mode='markers',
                        name=f"{junta} {dir_lbl} {cond_lbl}",
                        marker=dict(symbol=sym, size=10, color=col_,
                                    line=dict(width=1.5, color="black")),
                        showlegend=False,
                        hovertemplate=(
                            f"<b>{junta} — Dir {dir_lbl} — {cond_lbl}</b><br>"
                            f"f = {f_val:.1f} Hz ({f_val*60:.0f} cpm)<br>"
                            f"A = {u_val:.5f} mm<extra></extra>")
                    ))

        fig_blake.add_vrect(x0=f_excl_lo*60, x1=f_excl_hi*60,
                            fillcolor="orange", opacity=0.15, line_width=0)
        fig_blake.add_vline(x=f_op*60, line_dash="dash", line_color="red", opacity=0.8)
        
        fig_blake.update_layout(
            title="Blake Fig. 10-2 — Amplitud vs Frecuencia",
            height=850, font=dict(family="Arial"), 
            legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            margin=dict(l=60, r=40, t=60, b=60), 
            plot_bgcolor='rgba(0,0,0,0)', 
            paper_bgcolor='rgba(0,0,0,0)'
        )
        
        fig_blake.update_xaxes(
            title="Frecuencia (cpm)", type="log", range=[2, 4], 
            tickmode="array", tickvals=X_TICK_VALS, ticktext=X_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        fig_blake.update_yaxes(
            title="Amplitud peak (mm)", type="log", range=[-4.5, 0.6], 
            tickmode="array", tickvals=Y_TICK_VALS, ticktext=Y_TICK_TEXT,
            **grid_style, minor=minor_grid_style
        )
        
        st.plotly_chart(fig_blake, use_container_width=True, config={'toImageButtonOptions': {'format': 'png', 'filename': 'Grafico_Blake', 'scale': 2}})

# ══════════════════════════════════════════════════════════════════════════════
# TAB 5: DISTORSIÓN ANGULAR
# ══════════════════════════════════════════════════════════════════════════════
with tab_dist:
    st.header("Distorsión Angular Dinámica — Diagrama Vectorial")
    if not resultados:
        st.info("Cargue datos para ver el análisis de distorsión angular.")
    else:
        casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
        casos_sel_dist = st.multiselect("Seleccionar casos", casos_uniq, default=casos_uniq[:1])
        f_eval = st.number_input("Frecuencia de evaluación (Hz)", value=f_op, step=0.1, format="%.2f")

        if casos_sel_dist:
            F_op_V_d = F_op_V; F_op_H_d = F_op_H

            for caso in casos_sel_dist:
                st.subheader(f"Caso: {caso}")
                grupo = {k: v for k, v in resultados.items() if k[0] == caso}
                if not grupo: continue

                juntas_c = sorted(set(k[1] for k in grupo))
                dir_lbl_c = next(iter(grupo.values()))['dir']
                F_f = F_op_V_d if dir_lbl_c == 'Z' else F_op_H_d

                # Obtener fase y amplitud a f_eval
                reales, imags, mags, fases, labels = [], [], [], [], []
                for junta in juntas_c:
                    r = grupo.get((caso, junta))
                    if r is None: continue
                    freqs = np.array(r['freqs'])
                    idx   = np.argmin(np.abs(freqs - f_eval))
                    frf   = np.array(r['frf_mm'])
                    fase  = np.array(r['fase_deg'])
                    mag_v = frf[idx] * F_f / F_REF_N
                    phi   = np.radians(fase[idx])
                    reales.append(mag_v * np.cos(phi))
                    imags.append(mag_v * np.sin(phi))
                    mags.append(mag_v)
                    fases.append(fase[idx])
                    labels.append(junta)

                if not labels:
                    st.warning("Sin datos para este caso.")
                    continue

                col_v1, col_v2 = st.columns([2, 1])
                with col_v1:
                    fig_v = go.Figure()
                    colors_v = plt.cm.Set1.colors
                    for k, (re, im, junta) in enumerate(zip(reales, imags, labels)):
                        col_ = '#%02x%02x%02x' % tuple(int(x*255) for x in colors_v[k % len(colors_v)][:3])
                        fig_v.add_trace(go.Scatter(
                            x=[0, re], y=[0, im], mode='lines+markers',
                            name=junta, line=dict(color=col_, width=3),
                            marker=dict(symbol=['circle','arrow'], size=[5,14],
                                        color=col_, angleref='previous'),
                        ))
                        fig_v.add_annotation(
                            x=re*1.2, y=im*1.2, text=f"<b>{junta}</b><br>{fases[k]:.0f}°",
                            showarrow=False, font=dict(size=11, color=col_)
                        )
                    fig_v.add_hline(y=0, line_color="gray", line_width=0.5)
                    fig_v.add_vline(x=0, line_color="gray", line_width=0.5)
                    lim = max(abs(max(reales+imags, key=abs))*1.6, 1e-9)
                    fig_v.update_layout(
                        title=f"Vectores de desplazamiento — {caso} @ {f_eval:.1f} Hz",
                        xaxis=dict(title="Re (mm)", range=[-lim, lim], scaleanchor="y"),
                        yaxis=dict(title="Im (mm)", range=[-lim, lim]),
                        height=450, font=dict(family="Arial"),
                        showlegend=True
                    )
                    st.plotly_chart(fig_v, use_container_width=True)

                with col_v2:
                    st.markdown("**Tabla de vectores**")
                    df_v = pd.DataFrame({
                        'Junta': labels,
                        'Mag (mm)': [round(m, 5) for m in mags],
                        'Fase (°)': [round(f, 1) for f in fases],
                    })
                    st.dataframe(df_v, use_container_width=True)

                    # Diferencial de fases
                    if len(fases) >= 2:
                        st.markdown("**Diferenciales de fase:**")
                        for i in range(len(labels)-1):
                            df_i = abs(fases[i] - fases[i+1])
                            if df_i > 180: df_i = 360 - df_i
                            emoji = "🔴" if df_i > 90 else "🟡" if df_i > 45 else "🟢"
                            st.markdown(
                                f"{emoji} **{labels[i]}–{labels[i+1]}:** {df_i:.1f}°"
                            )

# ══════════════════════════════════════════════════════════════════════════════
# TAB 6: REPORTE
# ══════════════════════════════════════════════════════════════════════════════
with tab_report:
    st.header("Generación de Reporte")
    if not resultados:
        st.info("Cargue datos para generar el reporte.")
    else:
        col_rp1, col_rp2 = st.columns(2)
        with col_rp1:
            titulo_rep  = st.text_input("Título del reporte", "Auditoría Dinámica — Plataforma de Soporte")
            proyecto    = st.text_input("Proyecto / Referencia", "EST-P730_RA-DISEÑO")
            autor       = st.text_input("Elaborado por", "")
        with col_rp2:
            revisor     = st.text_input("Revisado por", "")
            revision    = st.text_input("Revisión", "A")
            incluir_figs = st.checkbox("Incluir gráficos en el reporte", value=True)

        if st.button("📄 Generar Reporte Word (.docx)", type="primary"):
            with st.spinner("Generando reporte..."):
                # ── Crear documento Word ──────────────────────────────────────
                doc = Document()
                sec = doc.sections[0]
                sec.page_width  = Cm(21)
                sec.page_height = Cm(29.7)
                sec.left_margin = sec.right_margin = Cm(2.5)
                sec.top_margin  = sec.bottom_margin = Cm(2.5)

                def add_h1(text):
                    p = doc.add_paragraph()
                    r = p.add_run(text)
                    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
                    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after  = Pt(6)
                    pf = p.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.5

                def add_para(text):
                    p = doc.add_paragraph()
                    r = p.add_run(text)
                    r.font.name = 'Arial'; r.font.size = Pt(11)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p.paragraph_format.line_spacing = 1.5

                # Portada
                doc.add_paragraph()
                p = doc.add_paragraph()
                r = p.add_run(titulo_rep.upper())
                r.font.name='Arial'; r.font.size=Pt(20); r.font.bold=True
                r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

                doc.add_paragraph()
                tabla_port = doc.add_table(rows=6, cols=2)
                tabla_port.style = 'Table Grid'
                datos_port = [
                    ("Proyecto", proyecto),
                    ("Elaborado por", autor),
                    ("Revisado por", revisor),
                    ("Revisión", revision),
                    ("Frecuencia de operación", f"{f_op} Hz ({f_op*60:.0f} cpm)"),
                    ("Zona de exclusión", f"{f_excl_lo:.2f} – {f_excl_hi:.2f} Hz"),
                ]
                for ri, (k, v) in enumerate(datos_port):
                    tabla_port.rows[ri].cells[0].text = k
                    tabla_port.rows[ri].cells[1].text = v
                    tabla_port.rows[ri].cells[0].paragraphs[0].runs[0].font.bold = True

                doc.add_page_break()

                # Resumen de diagnóstico
                add_h1("1. Resumen de Diagnóstico de Resonancia")
                if not df_res.empty:
                    t = doc.add_table(rows=1+len(df_res), cols=len(df_res.columns))
                    t.style = 'Table Grid'
                    for ci, col_n in enumerate(df_res.columns):
                        c = t.rows[0].cells[ci]
                        c.text = col_n
                        r = c.paragraphs[0].runs[0]
                        r.font.bold = True; r.font.size = Pt(9)
                        r.font.name = 'Arial'
                        tcp = c._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(q('w:val'),'clear'); shd.set(q('w:color'),'auto')
                        shd.set(q('w:fill'),'1B2A4A'); tcp.append(shd)
                        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                    for ri, row_data in df_res.iterrows():
                        for ci, val in enumerate(row_data):
                            c = t.rows[ri+1].cells[ci]
                            c.text = str(val)
                            r = c.paragraphs[0].runs[0]
                            r.font.size = Pt(9); r.font.name = 'Arial'

                # Clasificaciones
                add_h1("2. Clasificación de Amplitudes")
                add_para(
                    f"Frecuencia de operación = {f_op} Hz | "
                    f"F_op_V = {F_op_V} N/pata | F_op_H = {F_op_H} N/pata | "
                    f"F_rd_V = {F_rd_V} N/pata | F_rd_H = {F_rd_H} N/pata"
                )

                # Figuras (matplotlib)
                if incluir_figs and resultados:
                    add_h1("3. Gráficos")
                    casos_uniq = sorted(set(r['caso'] for r in resultados.values()))
                    juntas_uniq = sorted(set(r['junta'] for r in resultados.values()))
                    n_c = min(len(casos_uniq), 3)
                    n_r = math.ceil(len(casos_uniq) / n_c)

                    fig_mpl, axes = plt.subplots(n_r, n_c, figsize=(14, 4*n_r))
                    if n_r*n_c == 1: axes = np.array([[axes]])
                    elif n_r == 1 or n_c == 1: axes = axes.reshape(n_r, n_c)

                    for ci, caso in enumerate(casos_uniq):
                        ax = axes[ci//n_c][ci%n_c]
                        grupo = {k: v for k, v in resultados.items() if k[0] == caso}
                        for (c, junta), r in grupo.items():
                            col_f = color_junta(juntas_uniq, junta)
                            ax.plot(r['freqs'], r['frf_mm'],
                                    color=col_f[:3], lw=1.8, label=junta)
                            ax.plot(r['f_pk'], r['frf_pk'], 'o', color=col_f[:3], ms=6)
                        ax.axvspan(f_excl_lo, f_excl_hi, color='orange', alpha=0.15)
                        ax.axvline(f_op, color='red', ls='--', lw=1.5, alpha=0.7)
                        ax.set_title(caso, fontsize=10, fontweight='bold')
                        ax.set_xlabel('Frecuencia (Hz)', fontsize=9)
                        ax.set_ylabel('FRF (mm/Ton)', fontsize=9)
                        ax.legend(fontsize=7.5, ncol=2)
                        ax.grid(True, alpha=0.25, ls='--')
                        ax.set_xlim(df_raw['Freq'].min()*0.95, df_raw['Freq'].max()*1.05)
                        ax.set_ylim(bottom=0)

                    plt.tight_layout()
                    buf = io.BytesIO()
                    fig_mpl.savefig(buf, format='png', dpi=160, bbox_inches='tight')
                    buf.seek(0)
                    plt.close(fig_mpl)

                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_img.add_run()
                    run.add_picture(buf, width=Cm(15))

                # Guardar y descargar
                buf_doc = io.BytesIO()
                doc.save(buf_doc)
                buf_doc.seek(0)

                st.success("✅ Reporte generado exitosamente")
                st.download_button(
                    label="⬇️ Descargar Reporte (.docx)",
                    data=buf_doc,
                    file_name=f"Auditoria_Dinamica_{revision}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ── Footer ────────────────────────────────────────────────────────────────────
st.divider()
st.caption(
    "Analizador Dinámico — ACI 351.3R-04 · ISO 20816-3:2022 · Richart (1962) · Blake (1964) "
    "| Datos: SAP2000 v27.1.0 (Kgf, cm, C) | F_ref = 1 Ton = 9810 N"
)
